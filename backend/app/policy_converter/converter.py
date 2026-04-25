"""
Template Conversion Engine — applies an organization's internal template
to a generated policy/standard DOCX.

Operations performed (python-docx + OOXML manipulation):
  1. Clone the source policy DOCX as the working document.
  2. Inject org header / footer with logo, colors, left/center/right text.
  3. Override heading colors with org primary color.
  4. Replace cover block metadata (org name, logo, classification).
  5. Apply font family changes via document-level default styles.
  6. Update footer page field.
  7. Optionally apply template-DOCX base styles when a template file is provided.

Returns the path to the converted DOCX.
"""
import io
import os
import re
import copy
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_HEADER_FOOTER

from .schemas import ForensicDocumentMap, TemplateConfig


# ── Color utilities ───────────────────────────────────────────────────────────

def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    h = hex_color.lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


# ── Header / Footer builders ─────────────────────────────────────────────────

def _clear_header_footer(section, header: bool = True):
    """Remove all content from a header or footer."""
    hf = section.header if header else section.footer
    for p in hf.paragraphs:
        for run in p.runs:
            run.text = ""
    return hf


def _build_header(section, cfg: TemplateConfig, logo_bytes: Optional[bytes] = None):
    """Replace section header with org branding."""
    hdr = section.header
    hdr.is_linked_to_previous = False

    # Remove existing content
    for p in hdr.paragraphs:
        p.clear()

    hdr_p = hdr.paragraphs[0] if hdr.paragraphs else hdr.add_paragraph()
    hdr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r, g, b = _hex_to_rgb(cfg.primary_color)

    if logo_bytes:
        try:
            run_logo = hdr_p.add_run()
            run_logo.add_picture(io.BytesIO(logo_bytes), height=Cm(1.0))
            hdr_p.add_run("  ")
        except Exception:
            pass

    if cfg.header_left:
        run_l = hdr_p.add_run(cfg.header_left)
        run_l.font.color.rgb = RGBColor(r, g, b)
        run_l.font.size = Pt(9)
        hdr_p.add_run("  |  ").font.size = Pt(9)

    if cfg.header_center:
        run_c = hdr_p.add_run(cfg.header_center)
        run_c.font.color.rgb = RGBColor(r, g, b)
        run_c.font.size = Pt(9)
        hdr_p.add_run("  |  ").font.size = Pt(9)

    if cfg.header_right:
        run_r = hdr_p.add_run(cfg.header_right)
        run_r.bold = True
        run_r.font.size = Pt(9)
        run_r.font.color.rgb = RGBColor(r, g, b)

    # Bottom border line on header
    pPr = hdr_p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), cfg.primary_color.lstrip("#"))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _build_footer(section, cfg: TemplateConfig):
    """Replace section footer with org branding + page numbers."""
    ftr = section.footer
    ftr.is_linked_to_previous = False

    for p in ftr.paragraphs:
        p.clear()

    ftr_p = ftr.paragraphs[0] if ftr.paragraphs else ftr.add_paragraph()
    ftr_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r, g, b = _hex_to_rgb(cfg.secondary_color)

    if cfg.footer_left:
        ftr_p.add_run(cfg.footer_left + "   ").font.size = Pt(8)

    # Page number field
    if "{page}" in cfg.footer_center:
        before, after = cfg.footer_center.replace("{page}", "\x00").split("\x00", 1)
        if "{total}" in before:
            before = before.replace("{total}", "").strip()
        if "{total}" in after:
            after = after.replace("{total}", "").strip()

        ftr_p.add_run(before).font.size = Pt(8)

        # PAGE field
        run_pg = ftr_p.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        run_pg._r.append(fldChar)

        run_pg2 = ftr_p.add_run()
        instrText = OxmlElement("w:instrText")
        instrText.text = " PAGE "
        run_pg2._r.append(instrText)

        run_pg3 = ftr_p.add_run()
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run_pg3._r.append(fldChar2)

        if after:
            ftr_p.add_run(after).font.size = Pt(8)

    elif cfg.footer_center:
        ftr_p.add_run(cfg.footer_center).font.size = Pt(8)

    if cfg.footer_right:
        ftr_p.add_run("   " + cfg.footer_right).font.size = Pt(8)

    # Style the footer text
    for run in ftr_p.runs:
        run.font.color.rgb = RGBColor(r, g, b)


# ── Heading color override ────────────────────────────────────────────────────

def _recolor_headings(doc: Document, hex_color: str):
    """Apply org primary color to all heading paragraphs."""
    r, g, b = _hex_to_rgb(hex_color)
    for para in doc.paragraphs:
        if para.style and para.style.name.startswith("Heading"):
            for run in para.runs:
                run.font.color.rgb = RGBColor(r, g, b)


# ── Cover page update ─────────────────────────────────────────────────────────

def _update_cover(doc: Document, cfg: TemplateConfig, logo_bytes: Optional[bytes] = None):
    """
    Update cover paragraph: replace org name reference and inject logo.
    The cover is the first few paragraphs before the first page break.
    """
    r, g, b = _hex_to_rgb(cfg.primary_color)
    cover_done = False

    for i, para in enumerate(doc.paragraphs):
        if i > 15:  # Cover block is always within first 15 paragraphs
            break

        text = para.text
        # Replace any existing org name reference
        if text and any(w in text for w in ["Organization:", "Gulf National Bank", "org_name"]):
            for run in para.runs:
                if "Gulf National Bank" in run.text:
                    run.text = run.text.replace("Gulf National Bank", cfg.org_name)
                elif "Organization:" in run.text:
                    run.text = f"Organization: {cfg.org_name}"

        # Apply primary color to title runs
        if para.style and para.style.name.startswith("Heading") or (
            para.runs and any(r.bold for r in para.runs)
        ):
            if not cover_done:
                for run in para.runs:
                    if run.font.size and run.font.size >= Pt(16):
                        run.font.color.rgb = RGBColor(r, g, b)

        # Check for page break — end of cover
        if 'w:type="page"' in para._p.xml:
            cover_done = True

        # Inject logo after title block (first bold+large run area)
        if logo_bytes and i == 0 and not cover_done:
            try:
                logo_run = para.add_run()
                para.add_run("\n")
                logo_run.add_picture(io.BytesIO(logo_bytes), height=Cm(1.5))
            except Exception:
                pass


# ── Font family override ──────────────────────────────────────────────────────

def _set_default_font(doc: Document, font_name: str):
    """Override document default font via rPrDefault in styles."""
    styles_element = doc.styles.element
    docDefaults = styles_element.find(qn("w:docDefaults"))
    if docDefaults is None:
        return
    rPrDefault = docDefaults.find(qn("w:rPrDefault"))
    if rPrDefault is None:
        rPrDefault = OxmlElement("w:rPrDefault")
        docDefaults.append(rPrDefault)
    rPr = rPrDefault.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        rPrDefault.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), font_name)
    rFonts.set(qn("w:hAnsi"), font_name)


# ── Apply template DOCX styles ────────────────────────────────────────────────

def _copy_styles_from_template(source_doc: Document, template_doc: Document):
    """
    Copy heading and paragraph styles from a template DOCX into source_doc.
    Only copies styles that exist in both documents.
    """
    template_styles = {s.name: s for s in template_doc.styles}
    for style in source_doc.styles:
        if style.name in template_styles:
            tmpl_style = template_styles[style.name]
            # Copy font from template
            if tmpl_style.font.name:
                style.font.name = tmpl_style.font.name
            if tmpl_style.font.size:
                style.font.size = tmpl_style.font.size
            if tmpl_style.font.color.type:
                try:
                    style.font.color.rgb = tmpl_style.font.color.rgb
                except Exception:
                    pass


# ── Main conversion function ─────────────────────────────────────────────────

def convert_document(
    fmap: ForensicDocumentMap,
    cfg: TemplateConfig,
    output_dir: str,
    logo_path: Optional[str] = None,
    template_docx_path: Optional[str] = None,
) -> str:
    """
    Apply the org template to the source DOCX and return path to converted file.

    Args:
        fmap: Forensic map for the document.
        cfg: Template configuration (colors, fonts, header/footer text).
        output_dir: Where to write the converted DOCX.
        logo_path: Optional path to a PNG/JPEG logo file.
        template_docx_path: Optional org template DOCX for style inheritance.

    Returns:
        Path to the converted DOCX file.
    """
    os.makedirs(output_dir, exist_ok=True)
    src = fmap.source_docx
    out_filename = f"{fmap.doc_id}_{cfg.org_id}_converted.docx"
    out_path = os.path.join(output_dir, out_filename)

    # Clone source DOCX
    shutil.copy2(src, out_path)
    doc = Document(out_path)

    # Load logo bytes if provided
    logo_bytes = None
    if logo_path and os.path.exists(logo_path):
        with open(logo_path, "rb") as f:
            logo_bytes = f.read()

    # Load template DOCX if provided
    template_doc = None
    if template_docx_path and os.path.exists(template_docx_path):
        template_doc = Document(template_docx_path)

    # 1. Apply template styles (if template DOCX provided)
    if template_doc:
        _copy_styles_from_template(doc, template_doc)

    # 2. Default font
    if cfg.font_family:
        _set_default_font(doc, cfg.font_family)

    # 3. Heading colors
    _recolor_headings(doc, cfg.primary_color)

    # 4. Cover page update
    _update_cover(doc, cfg, logo_bytes)

    # 5. Header / Footer on all sections
    for section in doc.sections:
        _build_header(section, cfg, logo_bytes)
        _build_footer(section, cfg)

    # 6. Save
    doc.save(out_path)
    print(f"[Converter] Converted: {out_path}")
    return out_path


def update_forensic_after_conversion(
    fmap: ForensicDocumentMap,
    converted_docx_path: str,
    forensic_maps_dir: str,
) -> ForensicDocumentMap:
    """
    Re-run forensic extraction on the converted DOCX to update page numbers
    in the forensic map (converted_page_number fields).
    """
    import json
    from .forensics import _count_pages, _CONTROLS_RE, _ELEMENT_RE, _REQ_ID_RE

    records, max_page = _count_pages(converted_docx_path)
    n = len(records)

    is_policy = fmap.document_type == "policy"
    seen: dict[tuple, int] = {}  # (element_ref, ctrl_id) -> page_number

    for i, rec in enumerate(records):
        text = rec["text"]
        if not text:
            continue
        if is_policy:
            m = _ELEMENT_RE.match(text)
            if m:
                element_ref = f"Element {m.group(1)}"
                for j in range(i + 1, min(i + 5, n)):
                    cm = _CONTROLS_RE.match(records[j]["text"])
                    if cm:
                        for ctrl_id in [c.strip() for c in cm.group(1).split(",")]:
                            seen[(element_ref, ctrl_id)] = records[j]["page"]
                        break
        else:
            m = _REQ_ID_RE.match(text)
            if m:
                req_id = m.group(1)
                for j in range(i + 1, min(i + 5, n)):
                    cm = _CONTROLS_RE.match(records[j]["text"])
                    if cm:
                        for ctrl_id in [c.strip() for c in cm.group(1).split(",")]:
                            seen[(req_id, ctrl_id)] = records[j]["page"]
                        break

    for cl in fmap.control_locations:
        key = (cl.element_ref, cl.control_id)
        if key in seen:
            cl.converted_page_number = seen[key]

    for sl in fmap.section_locations:
        # Try matching in converted records
        for rec in records:
            if rec["is_heading"] and sl.title.split()[0] in rec["text"]:
                sl.converted_page_number = rec["page"]
                break

    fmap.converted_docx = converted_docx_path
    fmap.converted_at = datetime.now(timezone.utc).isoformat()

    # Persist updated map
    out_path = os.path.join(forensic_maps_dir, f"{fmap.doc_id}_forensic_map.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(fmap.model_dump(), f, indent=2)

    return fmap
