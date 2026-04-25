"""
Forensic DOCX Reader — extracts page numbers and paragraph locations
for every control citation in a generated policy or standard DOCX.

Strategy (heuristic — no COM automation needed):
  1. Parse DOCX XML via python-docx to get all paragraphs + tables.
  2. Count explicit page breaks (<w:br w:type="page"/>) to track page changes.
  3. Match paragraphs against known rendered patterns from renderer.py.
  4. Map control_id → page_number, paragraph_index.
  5. Write new ForensicDocumentMap JSON alongside source artifacts.

Accuracy: ±1 page for dense content, exact for page-break boundaries.
"""
import json
import glob
import os
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Optional

from docx import Document
from docx.oxml.ns import qn

from .schemas import (
    ForensicDocumentMap,
    ControlLocation,
    SectionLocation,
)


# ── Helpers ───────────────────────────────────────────────────────────────────

def _has_page_break(para) -> bool:
    """Return True if paragraph XML contains an explicit page break."""
    xml = para._p.xml
    return 'w:type="page"' in xml or 'w:pageBreak' in xml


def _heading_level(para) -> Optional[int]:
    """Return heading level (1-9) or None for non-heading paragraphs."""
    style = para.style.name if para.style else ""
    m = re.match(r"Heading (\d+)", style)
    return int(m.group(1)) if m else None


def _iter_doc_paragraphs(docx_path: str):
    """
    Yield (text, style_name, is_heading, heading_level, para_obj) for every
    paragraph including those inside tables (in document order).
    """
    doc = Document(docx_path)
    # Iterate over body XML children to preserve order of paragraphs + tables
    body = doc.element.body
    para_idx = 0
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            # Direct paragraph
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            level = _heading_level(p)
            yield {
                'text': p.text.strip(),
                'style': p.style.name if p.style else '',
                'is_heading': level is not None,
                'heading_level': level,
                'para_obj': p,
                'para_idx': para_idx,
                'in_table': False,
            }
            para_idx += 1
        elif tag == 'tbl':
            # Table — iterate cells row by row
            from docx.table import Table
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        level = _heading_level(p)
                        yield {
                            'text': p.text.strip(),
                            'style': p.style.name if p.style else '',
                            'is_heading': False,
                            'heading_level': None,
                            'para_obj': p,
                            'para_idx': para_idx,
                            'in_table': True,
                        }
                        para_idx += 1


def _count_pages(docx_path: str) -> tuple[list[dict], int]:
    """
    Build a list of all paragraph records with estimated page numbers.
    Returns (paragraphs, max_page).
    """
    doc = Document(docx_path)
    records = []
    page = 1
    para_idx = 0

    body = doc.element.body
    for child in body.iterchildren():
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'p':
            from docx.text.paragraph import Paragraph
            p = Paragraph(child, doc)
            # Check for page break BEFORE this paragraph's content
            xml = p._p.xml
            if 'w:type="page"' in xml:
                page += 1
            # pPr pageBreakBefore
            pPr = p._p.find(qn('w:pPr'))
            if pPr is not None:
                pb = pPr.find(qn('w:pageBreakBefore'))
                if pb is not None:
                    val = pb.get(qn('w:val'), 'true')
                    if val not in ('0', 'false'):
                        page += 1

            level = _heading_level(p)
            records.append({
                'text': p.text.strip(),
                'style': p.style.name if p.style else '',
                'is_heading': level is not None,
                'heading_level': level,
                'para_idx': para_idx,
                'page': page,
                'in_table': False,
            })
            para_idx += 1

        elif tag == 'tbl':
            from docx.table import Table
            tbl = Table(child, doc)
            for row in tbl.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        xml = p._p.xml
                        if 'w:type="page"' in xml:
                            page += 1
                        level = _heading_level(p)
                        records.append({
                            'text': p.text.strip(),
                            'style': p.style.name if p.style else '',
                            'is_heading': False,
                            'heading_level': None,
                            'para_idx': para_idx,
                            'page': page,
                            'in_table': True,
                        })
                        para_idx += 1

    return records, page


# ── Section heading patterns from renderer.py ──────────────────────────────

_SECTION_PATTERNS = [
    # Policy
    (re.compile(r"^(\d+)\.\s{1,3}Objectives"), "Objectives"),
    (re.compile(r"^(\d+)\.\s{1,3}Scope"), "Scope and Applicability"),
    (re.compile(r"^(\d+)\.\s{1,3}Policy Elements"), "Policy Elements"),
    (re.compile(r"^(\d+)\.\s{1,3}Roles"), "Roles and Responsibilities"),
    (re.compile(r"^(\d+)\.\s{1,3}Policy Compliance"), "Policy Compliance"),
    (re.compile(r"^(\d+)\.\s{1,3}Exceptions"), "Exceptions"),
    # Standard
    (re.compile(r"^(\d+)\.\s{1,3}Standards Requirements"), "Standards Requirements"),
    (re.compile(r"^(\d+)\.\s{1,3}Update"), "Update and Review"),
    (re.compile(r"^(\d+)\.\s{1,3}Compliance"), "Compliance with the Standard"),
    # Sub-heading
    (re.compile(r"^(3a)\.\s{1,3}Cybersecurity Policy Suite"), "Policy Suite Catalog"),
    # Procedure
    (re.compile(r"^(\d+)\.\s{1,3}Procedure Objective"), "Procedure Objective"),
    (re.compile(r"^(\d+)\.\s{1,3}Objective"), "Procedure Objective"),
    (re.compile(r"^(\d+)\.\s{1,3}Prerequisites"), "Triggers, Prerequisites, Inputs and Tools"),
    (re.compile(r"^(\d+)\.\s{1,3}Detailed Procedure"), "Detailed Procedure Steps"),
    (re.compile(r"^(\d+)\.\s{1,3}Verification"), "Verification and Testing"),
    (re.compile(r"^(\d+)\.\s{1,3}Evidence"), "Evidence Collection"),
    (re.compile(r"^(\d+)\.\s{1,3}Decision"), "Decision Points, Exceptions and Escalations"),
]

_ELEMENT_RE = re.compile(r"^Element\s+(\d+)\.\s+(.+)", re.DOTALL)
_CONTROLS_RE = re.compile(r"^Controls:\s+(.+)")
_REQ_ID_RE   = re.compile(r"^(\d+\.\d+)\s+(.+)")
_DOMAIN_RE   = re.compile(r"^Domain\s+(\d+):\s+(.+)")
_STEP_RE     = re.compile(r"^Step\s+(\d+(?:\.\d+)?)[:\s]", re.IGNORECASE)
_PHASE_RE    = re.compile(r"^Phase\s+\d+[:\s]|^(Preparation|Implementation|Verification|Testing)[:\s]", re.IGNORECASE)


# ── Main extraction function ───────────────────────────────────────────────

def extract_forensic_map(
    doc_id: str,
    docx_path: str,
    result_json_path: str,
    bundles_json_path: str,
    draft_json_path: str,
    traceability_docx_path: str,
    output_dir: str = "forensic_maps",
) -> ForensicDocumentMap:
    """
    Extract forensic location data from a generated DOCX and its JSON artifacts.
    Returns a ForensicDocumentMap and writes it to output_dir.
    """
    print(f"[Forensics] Extracting: {doc_id} from {os.path.basename(docx_path)}")

    # ── Load source artifacts ──────────────────────────────────────────────
    with open(result_json_path, encoding="utf-8") as f:
        result = json.load(f)
    with open(bundles_json_path, encoding="utf-8") as f:
        bundles = json.load(f)
    with open(draft_json_path, encoding="utf-8") as f:
        draft = json.load(f)

    doc_type = draft.get("document_type", "policy")
    is_procedure = doc_type == "procedure" or "phases" in draft
    is_policy = not is_procedure and (doc_type == "policy" or "policy_elements" in draft)

    # Bundle lookup by control_id
    bundle_map: dict[str, dict] = {}
    for b in bundles:
        bundle_map[b["control_id"]] = b
        # Also index by chunk_id for secondary lookup
        if "chunk_id" in b:
            bundle_map[b["chunk_id"]] = b

    # ── Scan DOCX paragraphs ───────────────────────────────────────────────
    records, max_page = _count_pages(docx_path)
    total_paras = len(records)
    print(f"[Forensics] {total_paras} paragraphs, ~{max_page} pages")

    # ── Extract section locations ──────────────────────────────────────────
    section_locations: list[SectionLocation] = []
    seen_sections: set[str] = set()

    for rec in records:
        text = rec["text"]
        if not rec["is_heading"] and not text:
            continue
        for pattern, canon_title in _SECTION_PATTERNS:
            m = pattern.match(text)
            if m:
                sec_no = m.group(1)
                if sec_no not in seen_sections:
                    seen_sections.add(sec_no)
                    section_locations.append(SectionLocation(
                        section_number=sec_no,
                        title=canon_title,
                        page_number=rec["page"],
                        paragraph_index=rec["para_idx"],
                        heading_level=rec.get("heading_level") or 1,
                    ))
                break

    # ── Build expected control map from draft JSON ─────────────────────────
    # Maps element_ref → list[control_id]
    expected: dict[str, list[str]] = {}

    if is_policy:
        for elem in draft.get("policy_elements", []):
            ref = f"Element {elem['element_no']}"
            cids = [t["control_id"] for t in elem.get("trace", [])]
            expected[ref] = cids
        section_label = "Policy Elements"
        section_no_str = "3"
    elif is_procedure:
        for phase in draft.get("phases", []):
            for step in phase.get("steps", []):
                ref = f"Step {step.get('step_no', '?')}"
                # citations are chunk_ids — resolve to control_ids via bundle_map
                cids = []
                for cid in step.get("citations", []):
                    b = bundle_map.get(cid, {})
                    ctrl_id = b.get("control_id", cid)
                    if ctrl_id:
                        cids.append(ctrl_id)
                expected[ref] = cids
        section_label = "Detailed Procedure Steps"
        section_no_str = "7"
    else:
        for dom in draft.get("domains", []):
            dom_no = dom["domain_number"]
            dom_title = dom["title_en"]
            for req in dom.get("requirements", []):
                ref = req["req_id"]
                cids = [t["control_id"] for t in req.get("trace", [])]
                expected[ref] = cids
        section_label = "Standards Requirements"
        section_no_str = "3"

    # ── Scan DOCX for element / req paragraphs and control citations ───────
    control_locations: list[ControlLocation] = []
    seen_control_keys: set[str] = set()   # (element_ref, control_id)

    n = len(records)
    for i, rec in enumerate(records):
        text = rec["text"]
        if not text:
            continue

        if is_policy:
            m = _ELEMENT_RE.match(text)
            if m:
                element_no = m.group(1)
                element_ref = f"Element {element_no}"
                # Scan next paragraphs for "Controls: ..." line
                for j in range(i + 1, min(i + 5, n)):
                    ctrl_text = records[j]["text"]
                    cm = _CONTROLS_RE.match(ctrl_text)
                    if cm:
                        ctrl_ids_raw = [c.strip() for c in cm.group(1).split(",")]
                        ctrl_para_page = records[j]["page"]
                        ctrl_para_idx  = records[j]["para_idx"]
                        for ctrl_id in ctrl_ids_raw:
                            ctrl_id = ctrl_id.strip()
                            if not ctrl_id:
                                continue
                            key = (element_ref, ctrl_id)
                            if key in seen_control_keys:
                                continue
                            seen_control_keys.add(key)
                            b = bundle_map.get(ctrl_id, {})
                            control_locations.append(ControlLocation(
                                control_id=ctrl_id,
                                framework=b.get("framework", "NCA_ECC"),
                                nca_id=b.get("nca_id", ctrl_id),
                                chunk_id=b.get("chunk_id"),
                                rerank_score=b.get("rerank_score"),
                                element_ref=element_ref,
                                section=section_label,
                                section_number=section_no_str,
                                page_number=ctrl_para_page,
                                paragraph_index=ctrl_para_idx,
                                extraction_method="heuristic",
                                extraction_confidence=0.80,
                            ))
                        break
        elif is_procedure:
            m = _STEP_RE.match(text)
            if m:
                step_no = m.group(1) if m.lastindex and m.group(1) else text.split()[1].rstrip(":")
                element_ref = f"Step {step_no}"
                # Scan next paragraphs for "Controls: ..." line or collect from draft
                for j in range(i + 1, min(i + 8, n)):
                    ctrl_text = records[j]["text"]
                    cm = _CONTROLS_RE.match(ctrl_text)
                    if cm:
                        ctrl_ids_raw = [c.strip() for c in cm.group(1).split(",")]
                        ctrl_para_page = records[j]["page"]
                        ctrl_para_idx  = records[j]["para_idx"]
                        for ctrl_id in ctrl_ids_raw:
                            ctrl_id = ctrl_id.strip()
                            if not ctrl_id:
                                continue
                            key = (element_ref, ctrl_id)
                            if key in seen_control_keys:
                                continue
                            seen_control_keys.add(key)
                            b = bundle_map.get(ctrl_id, {})
                            control_locations.append(ControlLocation(
                                control_id=ctrl_id,
                                framework=b.get("framework", "NCA_ECC"),
                                nca_id=b.get("nca_id", ctrl_id),
                                chunk_id=b.get("chunk_id"),
                                rerank_score=b.get("rerank_score"),
                                element_ref=element_ref,
                                section=section_label,
                                section_number=section_no_str,
                                page_number=records[i]["page"],
                                paragraph_index=records[i]["para_idx"],
                                extraction_method="heuristic",
                                extraction_confidence=0.75,
                            ))
                        break
        else:
            # Standard — req_id like "1.1", "2.3"
            m = _REQ_ID_RE.match(text)
            if m:
                req_id = m.group(1)
                # Find domain context
                dom_title = None
                for dom in draft.get("domains", []):
                    for req in dom.get("requirements", []):
                        if req["req_id"] == req_id:
                            dom_title = dom["title_en"]
                            break
                    if dom_title:
                        break
                # Scan next paragraphs for "Controls: ..." line
                for j in range(i + 1, min(i + 5, n)):
                    ctrl_text = records[j]["text"]
                    cm = _CONTROLS_RE.match(ctrl_text)
                    if cm:
                        ctrl_ids_raw = [c.strip() for c in cm.group(1).split(",")]
                        ctrl_para_page = records[j]["page"]
                        ctrl_para_idx  = records[j]["para_idx"]
                        for ctrl_id in ctrl_ids_raw:
                            ctrl_id = ctrl_id.strip()
                            if not ctrl_id:
                                continue
                            key = (req_id, ctrl_id)
                            if key in seen_control_keys:
                                continue
                            seen_control_keys.add(key)
                            b = bundle_map.get(ctrl_id, {})
                            control_locations.append(ControlLocation(
                                control_id=ctrl_id,
                                framework=b.get("framework", "NCA_ECC"),
                                nca_id=b.get("nca_id", ctrl_id),
                                chunk_id=b.get("chunk_id"),
                                rerank_score=b.get("rerank_score"),
                                element_ref=req_id,
                                section=section_label,
                                section_number=section_no_str,
                                domain_title=dom_title,
                                page_number=ctrl_para_page,
                                paragraph_index=ctrl_para_idx,
                                extraction_method="heuristic",
                                extraction_confidence=0.80,
                            ))
                        break

    # ── Fallback: map controls from draft trace where not found in DOCX ────
    found_keys = {(c.element_ref, c.control_id) for c in control_locations}

    if is_policy:
        for elem in draft.get("policy_elements", []):
            ref = f"Element {elem['element_no']}"
            for t in elem.get("trace", []):
                ctrl_id = t["control_id"]
                key = (ref, ctrl_id)
                if key not in found_keys:
                    b = bundle_map.get(ctrl_id, {})
                    pg = next(
                        (sl.page_number for sl in section_locations if sl.section_number == "3"),
                        None
                    )
                    control_locations.append(ControlLocation(
                        control_id=ctrl_id,
                        framework=t.get("framework", "NCA_ECC"),
                        nca_id=b.get("nca_id", ctrl_id),
                        chunk_id=b.get("chunk_id") or t.get("source_ref"),
                        rerank_score=b.get("rerank_score"),
                        element_ref=ref,
                        section=section_label,
                        section_number=section_no_str,
                        page_number=pg,
                        extraction_method="heuristic",
                        extraction_confidence=0.50,
                    ))
                    found_keys.add(key)
    elif is_procedure:
        # Fallback: resolve chunk_ids from step citations → control_ids via bundle_map
        pg = next(
            (sl.page_number for sl in section_locations if sl.section_number == section_no_str),
            None
        )
        for phase in draft.get("phases", []):
            for step in phase.get("steps", []):
                step_no = step.get("step_no", "?")
                ref = f"Step {step_no}"
                for chunk_id in step.get("citations", []):
                    b = bundle_map.get(chunk_id, {})
                    ctrl_id = b.get("control_id", chunk_id)
                    if not ctrl_id:
                        continue
                    key = (ref, ctrl_id)
                    if key not in found_keys:
                        control_locations.append(ControlLocation(
                            control_id=ctrl_id,
                            framework=b.get("framework", "NCA_ECC"),
                            nca_id=b.get("nca_id", ctrl_id),
                            chunk_id=chunk_id,
                            rerank_score=b.get("rerank_score"),
                            element_ref=ref,
                            section=section_label,
                            section_number=section_no_str,
                            page_number=pg,
                            extraction_method="heuristic",
                            extraction_confidence=0.50,
                        ))
                        found_keys.add(key)
    else:
        for dom in draft.get("domains", []):
            for req in dom.get("requirements", []):
                req_id = req["req_id"]
                for t in req.get("trace", []):
                    ctrl_id = t["control_id"]
                    key = (req_id, ctrl_id)
                    if key not in found_keys:
                        b = bundle_map.get(ctrl_id, {})
                        pg = next(
                            (sl.page_number for sl in section_locations if sl.section_number == "3"),
                            None
                        )
                        control_locations.append(ControlLocation(
                            control_id=ctrl_id,
                            framework=t.get("framework", "NCA_ECC"),
                            nca_id=b.get("nca_id", ctrl_id),
                            chunk_id=b.get("chunk_id") or t.get("source_ref"),
                            rerank_score=b.get("rerank_score"),
                            element_ref=req_id,
                            section=section_label,
                            section_number=section_no_str,
                            domain_title=dom.get("title_en"),
                            page_number=pg,
                            extraction_method="heuristic",
                            extraction_confidence=0.50,
                        ))
                        found_keys.add(key)

    print(f"[Forensics] Found {len(section_locations)} sections, "
          f"{len(control_locations)} control citations")

    # ── Assemble ForensicDocumentMap ───────────────────────────────────────
    fmap = ForensicDocumentMap(
        doc_id=doc_id,
        document_type=doc_type if doc_type in ("policy", "standard", "procedure") else "policy",
        title=result.get("title", draft.get("title_en", "")),
        version=result.get("version", draft.get("version", "1.0")),
        org_name=draft.get("org_name", "Gulf National Bank"),
        source_result_json=result_json_path,
        source_bundles_json=bundles_json_path,
        source_draft_json=draft_json_path,
        source_docx=docx_path,
        source_traceability_docx=traceability_docx_path,
        qa_passed=result.get("qa_passed", False),
        shall_count=result.get("shall_count", 0),
        traced_count=result.get("traced_count", 0),
        forensic_extracted_at=datetime.now(timezone.utc).isoformat(),
        forensic_method="heuristic",
        total_paragraphs=total_paras,
        estimated_pages=max_page,
        section_locations=section_locations,
        control_locations=control_locations,
    )

    # ── Write output ───────────────────────────────────────────────────────
    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{doc_id}_forensic_map.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(fmap.model_dump(), f, indent=2)
    print(f"[Forensics] Written: {out_path}")

    return fmap


def update_from_officejs(
    fmap: ForensicDocumentMap,
    officejs_data: dict,
    output_dir: str = "forensic_maps",
) -> ForensicDocumentMap:
    """
    Update a forensic map with accurate page numbers received from the Office.js add-in.
    officejs_data format:
      {
        "control_locations": [{"control_id": "1-3-2", "element_ref": "Element 1", "page_number": 4, ...}],
        "section_locations": [{"section_number": "3", "page_number": 4}]
      }
    """
    ctrl_updates = {
        (item["element_ref"], item["control_id"]): item
        for item in officejs_data.get("control_locations", [])
    }
    sect_updates = {
        item["section_number"]: item
        for item in officejs_data.get("section_locations", [])
    }

    for cl in fmap.control_locations:
        key = (cl.element_ref, cl.control_id)
        if key in ctrl_updates:
            upd = ctrl_updates[key]
            cl.page_number = upd.get("page_number", cl.page_number)
            cl.paragraph_index = upd.get("paragraph_index", cl.paragraph_index)
            cl.extraction_method = "officejs"
            cl.extraction_confidence = 1.0

    for sl in fmap.section_locations:
        if sl.section_number in sect_updates:
            upd = sect_updates[sl.section_number]
            sl.page_number = upd.get("page_number", sl.page_number)
            sl.converted_page_number = upd.get("converted_page_number")

    fmap.forensic_method = "officejs"
    fmap.forensic_extracted_at = datetime.now(timezone.utc).isoformat()

    os.makedirs(output_dir, exist_ok=True)
    out_path = os.path.join(output_dir, f"{fmap.doc_id}_forensic_map.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(fmap.model_dump(), f, indent=2)

    return fmap


# ── Discovery helpers ─────────────────────────────────────────────────────────

def find_artifacts(policy_output_dir: str, doc_id: str) -> dict:
    """
    Locate the latest artifact JSON files for a given doc_id.
    Returns dict with keys: result_json, bundles_json, draft_json,
    main_docx, traceability_docx.
    Raises FileNotFoundError if essential files are missing.
    """
    glob_result = sorted(glob.glob(
        os.path.join(policy_output_dir, f"*_{doc_id.replace('-', '_')}*_0_result.json")
    ))
    # doc_id is like POL-01; search by doc_id in result JSON content
    results_all = sorted(glob.glob(os.path.join(policy_output_dir, "*_0_result.json")))
    result_path = None
    for rp in reversed(results_all):  # latest first
        try:
            with open(rp) as f:
                d = json.load(f)
            if d.get("doc_id") == doc_id:
                result_path = rp
                break
        except Exception:
            continue

    if not result_path:
        raise FileNotFoundError(f"No result JSON found for {doc_id}")

    prefix = result_path.replace("_0_result.json", "")
    bundles = prefix + "_1_bundles.json"
    drafts = sorted(glob.glob(prefix + "_2_draft_attempt*.json"))
    draft = drafts[-1] if drafts else None
    main_docx = os.path.join(policy_output_dir, f"{doc_id}.docx")
    trace_docx = os.path.join(policy_output_dir, f"{doc_id}_traceability.docx")

    for p, label in [(bundles, "bundles"), (main_docx, "main_docx")]:
        if not os.path.exists(p):
            raise FileNotFoundError(f"Missing {label}: {p}")

    return {
        "result_json": result_path,
        "bundles_json": bundles,
        "draft_json": draft,
        "main_docx": main_docx,
        "traceability_docx": trace_docx if os.path.exists(trace_docx) else "",
    }


def extract_all(
    policy_output_dir: str,
    forensic_maps_dir: str,
    doc_ids: Optional[list[str]] = None,
) -> list[ForensicDocumentMap]:
    """Extract forensic maps for all (or specified) documents."""
    if doc_ids is None:
        # Discover all doc_ids from result JSONs
        results = glob.glob(os.path.join(policy_output_dir, "*_0_result.json"))
        discovered = {}
        for rp in results:
            try:
                with open(rp) as f:
                    d = json.load(f)
                did = d.get("doc_id")
                if did and did not in discovered:
                    discovered[did] = rp
            except Exception:
                continue
        doc_ids = list(discovered.keys())

    maps = []
    for doc_id in sorted(doc_ids):
        try:
            arts = find_artifacts(policy_output_dir, doc_id)
            if not arts["draft_json"]:
                print(f"[Forensics] SKIP {doc_id}: no draft JSON found")
                continue
            fmap = extract_forensic_map(
                doc_id=doc_id,
                docx_path=arts["main_docx"],
                result_json_path=arts["result_json"],
                bundles_json_path=arts["bundles_json"],
                draft_json_path=arts["draft_json"],
                traceability_docx_path=arts["traceability_docx"],
                output_dir=forensic_maps_dir,
            )
            maps.append(fmap)
        except Exception as e:
            print(f"[Forensics] ERROR {doc_id}: {e}")
    return maps
