"""
Ministry Renderer Pro — professional DOCX renderer for both English (LTR)
and Arabic (RTL) ministry documents.

Uses python-docx + raw OOXML for:
  - Full-page branded cover page (dark green + white layout)
  - Internal page header (ministry name | doc ref)
  - Footer (classification | page number)
  - DRAFT watermark on internal pages
  - Professional section headings with colored rule
  - Alternating-row tables
"""
from __future__ import annotations
import os
from copy import deepcopy
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from .ministry_models import (
    MinistryPolicyDraft, MinistryStandardDraft, MinistryProcedureDraft,
    DefinitionEntry, ApprovalStage, VersionRow,
    StandardDomainCluster, ProcedurePhase, ProcedureRoleItem,
)

# ── Brand colours ──────────────────────────────────────────────────────────────
C_GREEN       = "0F4D43"
C_GREEN_LIGHT = "2E7D6B"
C_NEUTRAL     = "F4F4F4"
C_BORDER      = "D9D9D9"
C_ROW_ALT     = "EFF5F4"   # very light green for alternating rows
C_WHITE       = "FFFFFF"
C_TEXT        = "1A1A1A"
C_SUBHEADING  = "1D6B5E"

ORG_NAME = "Ministry of Communications and Information Technology"
ORG_NAME_SHORT = "MCIT"


# ═══════════════════════════════════════════════════════════════════════════════
# Low-level OOXML helpers
# ═══════════════════════════════════════════════════════════════════════════════

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # remove existing shd
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_cell_height(row, height_cm: float):
    """Set exact row height."""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    trHeight = OxmlElement("w:trHeight")
    trHeight.set(qn("w:val"), str(int(height_cm * 567)))  # cm -> twips (1cm = 567 twips)
    trHeight.set(qn("w:hRule"), "exact")
    trPr.append(trHeight)


def _set_table_width_full(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr")) or OxmlElement("w:tblPr")
    tblW = OxmlElement("w:tblW")
    tblW.set(qn("w:w"), "9898")  # ~full page width in 50ths of a percent
    tblW.set(qn("w:type"), "pct")
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    tblPr.append(tblW)
    # indent 0
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")
    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    tblPr.append(tblInd)


def _remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tblBorders.append(border)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _set_thin_borders(table):
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), C_BORDER)
        tblBorders.append(border)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblPr.append(tblBorders)


def _set_para_spacing(para, before_pt: int = 0, after_pt: int = 6, line_240: int = 276):
    """Set paragraph spacing. line_240 is in 240ths of a line (276 ≈ 1.15)."""
    pPr = para._p.get_or_add_pPr()
    spacing = pPr.find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        pPr.append(spacing)
    spacing.set(qn("w:before"), str(before_pt * 20))
    spacing.set(qn("w:after"), str(after_pt * 20))
    spacing.set(qn("w:line"), str(line_240))
    spacing.set(qn("w:lineRule"), "auto")


def _set_rtl(para):
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    pPr.append(jc)


def _set_run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)


def _para_align(para, align: str):
    """align: left | center | right | both"""
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), align)
    pPr.append(jc)


def _add_run(para, text: str, bold=False, italic=False,
             size_pt=11, color_hex=None, font="Calibri") -> None:
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size_pt)
    if color_hex:
        r, g, b = int(color_hex[:2], 16), int(color_hex[2:4], 16), int(color_hex[4:], 16)
        run.font.color.rgb = RGBColor(r, g, b)
    return run


def _page_number_field(para):
    """Append PAGE field code to paragraph."""
    run = para.add_run()
    for tag, text in [("begin", None), ("instrText", " PAGE "), ("end", None)]:
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)


def _numpages_field(para):
    run = para.add_run()
    for tag, text in [("begin", None), ("instrText", " NUMPAGES "), ("end", None)]:
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        else:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), tag)
        run._r.append(el)


# ═══════════════════════════════════════════════════════════════════════════════
# Watermark  (VML-based, same as native Word "Draft" watermark)
# ═══════════════════════════════════════════════════════════════════════════════

_WATERMARK_VML = """
<v:shape xmlns:v="urn:schemas-microsoft-com:vml"
         xmlns:o="urn:schemas-microsoft-com:office:office"
         xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
         id="PowerPlusWaterMarkObject" o:spid="_x0000_s2051"
         type="#_x0000_t136"
         style="position:absolute;margin-left:0;margin-top:0;width:527.85pt;height:131.95pt;z-index:-251655168;mso-position-horizontal:center;mso-position-horizontal-relative:margin;mso-position-vertical:center;mso-position-vertical-relative:margin"
         fillcolor="#C0C0C0" stroked="f">
  <v:textpath style="font-family:&quot;Calibri&quot;;font-size:1pt;font-weight:bold" string="DRAFT"
              trim="t" on="t"/>
</v:shape>
"""

def _add_watermark_to_header(header):
    """Add DRAFT watermark VML to a header paragraph."""
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    noProof = OxmlElement("w:noProof")
    rPr.append(noProof)
    r.append(rPr)
    pict = OxmlElement("w:pict")
    pict.append(etree.fromstring(_WATERMARK_VML))
    r.append(pict)
    para._p.append(r)


# ═══════════════════════════════════════════════════════════════════════════════
# Cover page builder
# ═══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc: Document, title: str, doc_id: str, doc_type: str,
                 version: str, classification: str, rtl: bool = False):
    """Build a full-page branded cover using a borderless table."""

    cover_table = doc.add_table(rows=3, cols=1)
    _remove_table_borders(cover_table)
    _set_table_width_full(cover_table)
    cover_table.style = "Table Grid"

    # ── Row 1: Dark green top band (org name + document type) ──
    top_row = cover_table.rows[0]
    _set_cell_height(top_row, 9.5)
    _set_cell_bg(top_row.cells[0], C_GREEN)
    top_cell = top_row.cells[0]
    top_cell._tc.get_or_add_tcPr()

    # Clear default paragraph
    top_cell.paragraphs[0].clear()
    # Spacer
    sp1 = top_cell.add_paragraph()
    _set_para_spacing(sp1, 0, 0)

    # Ministry name
    org_para = top_cell.add_paragraph()
    _para_align(org_para, "center")
    _set_para_spacing(org_para, 40, 4)
    _add_run(org_para, ORG_NAME, bold=True, size_pt=14, color_hex=C_WHITE, font="Calibri")

    # Doc type badge line
    type_para = top_cell.add_paragraph()
    _para_align(type_para, "center")
    _set_para_spacing(type_para, 8, 8)
    type_label = {"policy": "CYBERSECURITY POLICY", "standard": "CYBERSECURITY STANDARD", "procedure": "CYBERSECURITY PROCEDURE"}.get(doc_type, "DOCUMENT")
    _add_run(type_para, type_label, bold=True, size_pt=11, color_hex="A8D8CF", font="Calibri")

    # Divider line (using underline run trick)
    div_para = top_cell.add_paragraph()
    _para_align(div_para, "center")
    _set_para_spacing(div_para, 10, 0)
    div_run = div_para.add_run("─" * 40)
    div_run.font.color.rgb = RGBColor(0x4D, 0xA8, 0x8E)
    div_run.font.size = Pt(10)

    # ── Row 2: White title area ──
    mid_row = cover_table.rows[1]
    _set_cell_height(mid_row, 12.0)
    _set_cell_bg(mid_row.cells[0], C_WHITE)
    mid_cell = mid_row.cells[0]
    mid_cell.paragraphs[0].clear()

    # Generous top padding
    pad_para = mid_cell.add_paragraph()
    _set_para_spacing(pad_para, 30, 0)

    # Document ID
    id_para = mid_cell.add_paragraph()
    _para_align(id_para, "center")
    _set_para_spacing(id_para, 0, 6)
    _add_run(id_para, doc_id, bold=False, size_pt=14, color_hex=C_GREEN_LIGHT, font="Calibri")

    # Document title (large)
    title_para = mid_cell.add_paragraph()
    _para_align(title_para, "center")
    _set_para_spacing(title_para, 12, 12)
    _add_run(title_para, title, bold=True, size_pt=24, color_hex=C_GREEN, font="Calibri")

    # Decorative green rule below title
    rule_para = mid_cell.add_paragraph()
    _para_align(rule_para, "center")
    _set_para_spacing(rule_para, 8, 8)
    rule_run = rule_para.add_run()
    pPr = rule_para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom_bdr = OxmlElement("w:bottom")
    bottom_bdr.set(qn("w:val"), "single")
    bottom_bdr.set(qn("w:sz"), "12")
    bottom_bdr.set(qn("w:space"), "1")
    bottom_bdr.set(qn("w:color"), C_GREEN)
    pBdr.append(bottom_bdr)
    pPr.append(pBdr)

    # ── Row 3: Dark green footer strip (metadata) ──
    bot_row = cover_table.rows[2]
    _set_cell_height(bot_row, 5.0)
    _set_cell_bg(bot_row.cells[0], C_GREEN)
    bot_cell = bot_row.cells[0]
    bot_cell.paragraphs[0].clear()

    sp2 = bot_cell.add_paragraph()
    _set_para_spacing(sp2, 12, 0)

    # Metadata row: Version | Date | Classification
    meta_para = bot_cell.add_paragraph()
    _para_align(meta_para, "center")
    _set_para_spacing(meta_para, 4, 4)
    _add_run(meta_para, f"Version {version}   |   Issue Date: [Date of Approval]   |   {classification}",
             bold=False, size_pt=10, color_hex=C_WHITE, font="Calibri")

    doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# Header / Footer setup
# ═══════════════════════════════════════════════════════════════════════════════

def _setup_header_footer(doc: Document, doc_id: str, title: str,
                         classification: str, rtl: bool = False):
    """Configure headers and footers for all sections. Cover page gets blank."""
    for section in doc.sections:
        section.different_first_page_header_footer = True

        # ── First page header: blank (cover has its own content) ──
        fph = section.first_page_header
        fph.paragraphs[0].clear()

        # ── First page footer: blank ──
        fpf = section.first_page_footer
        fpf.paragraphs[0].clear()

        # ── Default header (all internal pages) ──
        hdr = section.header
        hdr.is_linked_to_previous = False
        if hdr.paragraphs:
            hdr.paragraphs[0].clear()
        else:
            hdr.add_paragraph()
        hp = hdr.paragraphs[0]

        # Add DRAFT watermark to header
        _add_watermark_to_header(hdr)

        # Header table: org name on left, doc ref on right
        hdr_table = hdr.add_table(rows=1, cols=2, width=Inches(6.3))
        _remove_table_borders(hdr_table)
        _set_table_width_full(hdr_table)

        lc = hdr_table.rows[0].cells[0]
        lc.paragraphs[0].clear()
        lp = lc.paragraphs[0]
        _set_para_spacing(lp, 0, 0)
        _add_run(lp, ORG_NAME_SHORT + " | Cybersecurity",
                 bold=True, size_pt=8, color_hex=C_GREEN, font="Calibri")

        rc = hdr_table.rows[0].cells[1]
        rc.paragraphs[0].clear()
        rp = rc.paragraphs[0]
        _para_align(rp, "right")
        _set_para_spacing(rp, 0, 0)
        _add_run(rp, f"{doc_id}  |  Version 1.0  |  {classification}",
                 bold=False, size_pt=8, color_hex=C_GREEN_LIGHT, font="Calibri")

        # Thin green bottom border on header
        pPr = hp._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "8")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_GREEN)
        pBdr.append(bot)
        pPr.append(pBdr)

        # ── Default footer ──
        ftr = section.footer
        ftr.is_linked_to_previous = False
        if ftr.paragraphs:
            ftr.paragraphs[0].clear()
        else:
            ftr.add_paragraph()

        ft = ftr.add_table(rows=1, cols=3, width=Inches(6.3))
        _remove_table_borders(ft)
        _set_table_width_full(ft)

        # Left: classification
        fl = ft.rows[0].cells[0]
        fl.paragraphs[0].clear()
        flp = fl.paragraphs[0]
        _set_para_spacing(flp, 0, 0)
        _add_run(flp, classification, bold=True, size_pt=8,
                 color_hex=C_GREEN, font="Calibri")

        # Center: page X / Y
        fc = ft.rows[0].cells[1]
        fc.paragraphs[0].clear()
        fcp = fc.paragraphs[0]
        _para_align(fcp, "center")
        _set_para_spacing(fcp, 0, 0)
        _add_run(fcp, "Page ", bold=False, size_pt=8, color_hex=C_TEXT, font="Calibri")
        _page_number_field(fcp)
        _add_run(fcp, " of ", bold=False, size_pt=8, color_hex=C_TEXT, font="Calibri")
        _numpages_field(fcp)

        # Right: doc ID
        fr = ft.rows[0].cells[2]
        fr.paragraphs[0].clear()
        frp = fr.paragraphs[0]
        _para_align(frp, "right")
        _set_para_spacing(frp, 0, 0)
        _add_run(frp, doc_id, bold=False, size_pt=8, color_hex=C_GREEN_LIGHT, font="Calibri")

        # Top border on footer
        fp0 = ftr.paragraphs[0]
        pPr2 = fp0._p.get_or_add_pPr()
        pBdr2 = OxmlElement("w:pBdr")
        top_bdr = OxmlElement("w:top")
        top_bdr.set(qn("w:val"), "single")
        top_bdr.set(qn("w:sz"), "8")
        top_bdr.set(qn("w:space"), "1")
        top_bdr.set(qn("w:color"), C_GREEN)
        pBdr2.append(top_bdr)
        pPr2.append(pBdr2)


# ═══════════════════════════════════════════════════════════════════════════════
# Section heading
# ═══════════════════════════════════════════════════════════════════════════════

def _section_heading(doc: Document, text: str, level: int = 1, rtl: bool = False):
    para = doc.add_paragraph()
    if rtl:
        _set_rtl(para)
    if level == 1:
        _set_para_spacing(para, before_pt=16, after_pt=4)
        _add_run(para, text, bold=True, size_pt=13, color_hex=C_GREEN, font="Calibri")
        # Bottom rule
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single")
        bot.set(qn("w:sz"), "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_GREEN_LIGHT)
        pBdr.append(bot)
        pPr.append(pBdr)
    else:
        _set_para_spacing(para, before_pt=10, after_pt=2)
        _add_run(para, text, bold=True, size_pt=11, color_hex=C_SUBHEADING, font="Calibri")
    return para


def _body_para(doc: Document, text: str, rtl: bool = False, indent_cm: float = 0):
    para = doc.add_paragraph()
    if rtl:
        _set_rtl(para)
    else:
        _para_align(para, "both")
    _set_para_spacing(para, before_pt=0, after_pt=6)
    if indent_cm:
        pPr = para._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(int(indent_cm * 567)))
        pPr.append(ind)
    _add_run(para, text, size_pt=11, font="Calibri")
    return para


def _bullet_para(doc: Document, text: str, rtl: bool = False):
    para = doc.add_paragraph()
    if rtl:
        _set_rtl(para)
    _set_para_spacing(para, before_pt=0, after_pt=3)
    pPr = para._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:hanging"), "180")
    pPr.append(ind)
    _add_run(para, "• " + text, size_pt=11, font="Calibri")
    return para


# ═══════════════════════════════════════════════════════════════════════════════
# Table builders
# ═══════════════════════════════════════════════════════════════════════════════

def _make_header_row(table, headers: list[str], rtl: bool = False):
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        _set_cell_bg(cell, C_GREEN)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        if rtl:
            _set_rtl(p)
        else:
            _para_align(p, "center")
        _set_para_spacing(p, 4, 4)
        text = headers[i] if i < len(headers) else ""
        _add_run(p, text, bold=True, size_pt=9, color_hex=C_WHITE, font="Calibri")


def _fill_data_row(row, texts: list[str], alt: bool = False, rtl: bool = False):
    bg = C_ROW_ALT if alt else C_WHITE
    for i, cell in enumerate(row.cells):
        _set_cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        if rtl:
            _set_rtl(p)
        _set_para_spacing(p, 2, 2)
        text = texts[i] if i < len(texts) else ""
        _add_run(p, text, size_pt=9, font="Calibri")


def _definitions_table(doc: Document, defs: list[DefinitionEntry], rtl: bool = False):
    if not defs:
        return
    headers = ["Definition", "Term"] if rtl else ["Term", "Definition"]
    table = doc.add_table(rows=1 + len(defs), cols=2)
    table.style = "Table Grid"
    _set_thin_borders(table)
    _set_table_width_full(table)
    _make_header_row(table, headers, rtl=rtl)
    for i, d in enumerate(defs):
        term_text = d.term_ar + (f" ({d.term_en})" if d.term_en and d.term_en != d.term_ar else "")
        def_text = d.definition_ar
        texts = [def_text, term_text] if rtl else [term_text, def_text]
        _fill_data_row(table.rows[i + 1], texts, alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()


def _approval_table(doc: Document, stages: list[ApprovalStage], rtl: bool = False):
    if not stages:
        return
    headers = ["Date", "Name", "Role / Title", "Stage"]
    table = doc.add_table(rows=1 + len(stages), cols=4)
    table.style = "Table Grid"
    _set_thin_borders(table)
    _set_table_width_full(table)
    _make_header_row(table, headers, rtl=rtl)
    for i, s in enumerate(stages):
        texts = [s.date_ar, s.name_ar, s.role_ar, s.stage_ar]
        _fill_data_row(table.rows[i + 1], texts, alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()


def _version_table(doc: Document, rows: list[VersionRow], rtl: bool = False):
    if not rows:
        return
    headers = ["Approval Date", "Updated By", "Summary", "Update Type", "Version"]
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    _set_thin_borders(table)
    _set_table_width_full(table)
    _make_header_row(table, headers, rtl=rtl)
    for i, r in enumerate(rows):
        texts = [r.approval_date, r.updated_by_ar, r.summary_ar, r.update_type_ar, r.version]
        _fill_data_row(table.rows[i + 1], texts, alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()


def _info_table(doc: Document, rows: list[tuple[str, str]], rtl: bool = False):
    table = doc.add_table(rows=len(rows), cols=2)
    table.style = "Table Grid"
    _set_thin_borders(table)
    _set_table_width_full(table)
    for i, (label, value) in enumerate(rows):
        lc = table.rows[i].cells[1 if rtl else 0]
        vc = table.rows[i].cells[0 if rtl else 1]
        _set_cell_bg(lc, C_NEUTRAL)
        lc.paragraphs[0].clear()
        lp = lc.paragraphs[0]
        if rtl:
            _set_rtl(lp)
        _set_para_spacing(lp, 2, 2)
        _add_run(lp, label, bold=True, size_pt=9, font="Calibri")
        vc.paragraphs[0].clear()
        vp = vc.paragraphs[0]
        if rtl:
            _set_rtl(vp)
        _set_para_spacing(vp, 2, 2)
        _add_run(vp, value, size_pt=9, font="Calibri")
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# New document factory
# ═══════════════════════════════════════════════════════════════════════════════

def _new_doc(rtl: bool = False) -> Document:
    doc = Document()
    # Default style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(2.54)
        section.right_margin  = Cm(2.54)
    return doc


# ═══════════════════════════════════════════════════════════════════════════════
# POLICY — Professional Renderer
# ═══════════════════════════════════════════════════════════════════════════════

def render_pro_policy(draft: MinistryPolicyDraft, output_dir: str,
                      lang: str = "en") -> str:
    """
    Render a MinistryPolicyDraft to a professional DOCX.
    lang='en' → LTR English layout
    lang='ar' → RTL Arabic layout
    """
    os.makedirs(output_dir, exist_ok=True)
    rtl = (lang == "ar")
    doc = _new_doc(rtl=rtl)

    title = draft.meta.title_en if lang == "en" else draft.meta.title_ar
    _build_cover(doc, title, draft.meta.doc_id, "policy",
                 draft.meta.version, draft.meta.classification, rtl=rtl)
    _setup_header_footer(doc, draft.meta.doc_id, title,
                         draft.meta.classification, rtl=rtl)

    # ── Document Information ──────────────────────────────────────────────────
    _section_heading(doc, "Document Information", rtl=rtl)
    info_rows = [
        ("Document Title", title),
        ("Document ID", draft.meta.doc_id),
        ("Document Type", "Cybersecurity Policy"),
        ("Version", draft.meta.version),
        ("Owner", draft.meta.owner),
        ("Classification", draft.meta.classification),
        ("Issue Date", draft.meta.issue_date or "[Date of Approval]"),
    ]
    _info_table(doc, info_rows, rtl=rtl)

    # ── Table of Contents ─────────────────────────────────────────────────────
    _section_heading(doc, "Table of Contents", rtl=rtl)
    toc_items = ["1.  Definitions", "2.  Policy Objective", "3.  Policy Scope",
                 "4.  Policy Statement", "5.  Roles and Responsibilities",
                 "6.  Related Documents", "7.  Effective Date",
                 "8.  Policy Review", "9.  Review and Approval", "10. Version Control"]
    for item in toc_items:
        _bullet_para(doc, item, rtl=rtl)
    doc.add_page_break()

    # ── Section 1: Definitions ────────────────────────────────────────────────
    _section_heading(doc, "1.  Definitions", rtl=rtl)
    _definitions_table(doc, draft.definitions, rtl=rtl)

    # ── Section 2: Objective ──────────────────────────────────────────────────
    _section_heading(doc, "2.  Policy Objective", rtl=rtl)
    _body_para(doc, draft.objective_ar, rtl=rtl)
    doc.add_paragraph()

    # ── Section 3: Scope ──────────────────────────────────────────────────────
    _section_heading(doc, "3.  Policy Scope", rtl=rtl)
    _body_para(doc, draft.scope_ar, rtl=rtl)
    doc.add_paragraph()

    # ── Section 4: Policy Statement (dominant) ────────────────────────────────
    _section_heading(doc, "4.  Policy Statement", rtl=rtl)
    _section_heading(doc, "General Provisions", level=2, rtl=rtl)
    for clause in draft.policy_clauses:
        p = doc.add_paragraph()
        if rtl:
            _set_rtl(p)
        else:
            _para_align(p, "both")
        _set_para_spacing(p, before_pt=2, after_pt=4)
        pPr = p._p.get_or_add_pPr()
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), "360")
        ind.set(qn("w:hanging"), "360")
        pPr.append(ind)
        num_run = p.add_run(f"{clause.clause_no}.  ")
        num_run.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
        num_run.font.name = "Calibri"
        txt_run = p.add_run(clause.text_ar)
        txt_run.font.size = Pt(11)
        txt_run.font.name = "Calibri"
        for bullet in clause.sub_bullets_ar:
            _bullet_para(doc, bullet, rtl=rtl)
    doc.add_paragraph()

    # ── Section 5: Roles ──────────────────────────────────────────────────────
    _section_heading(doc, "5.  Roles and Responsibilities", rtl=rtl)
    _body_para(doc, draft.roles_ar, rtl=rtl)
    doc.add_paragraph()

    # ── Section 6: Related Documents ─────────────────────────────────────────
    _section_heading(doc, "6.  Related Documents", rtl=rtl)
    if draft.related_docs:
        rel_table = doc.add_table(rows=1 + len(draft.related_docs), cols=2)
        rel_table.style = "Table Grid"
        _set_thin_borders(rel_table)
        _set_table_width_full(rel_table)
        _make_header_row(rel_table, ["Document Title", "Type"], rtl=rtl)
        for i, rd in enumerate(draft.related_docs):
            _fill_data_row(rel_table.rows[i + 1], [rd.title_ar, rd.ref_type], alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()

    # ── Section 7: Effective Date ─────────────────────────────────────────────
    _section_heading(doc, "7.  Effective Date", rtl=rtl)
    _body_para(doc, draft.effective_date_ar, rtl=rtl)
    doc.add_paragraph()

    # ── Section 8: Policy Review ──────────────────────────────────────────────
    _section_heading(doc, "8.  Policy Review", rtl=rtl)
    _body_para(doc, draft.review_ar, rtl=rtl)
    doc.add_paragraph()

    # ── Section 9: Review and Approval ───────────────────────────────────────
    _section_heading(doc, "9.  Review and Approval", rtl=rtl)
    _approval_table(doc, draft.approval_stages, rtl=rtl)

    # ── Section 10: Version Control ───────────────────────────────────────────
    _section_heading(doc, "10. Version Control", rtl=rtl)
    _version_table(doc, draft.version_rows, rtl=rtl)

    suffix = "_en" if lang == "en" else "_ar"
    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}{suffix}.docx")
    doc.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# STANDARD — Professional Renderer
# ═══════════════════════════════════════════════════════════════════════════════

def render_pro_standard(draft: MinistryStandardDraft, output_dir: str,
                        lang: str = "en") -> str:
    os.makedirs(output_dir, exist_ok=True)
    rtl = (lang == "ar")
    doc = _new_doc(rtl=rtl)
    title = draft.meta.title_en if lang == "en" else draft.meta.title_ar

    _build_cover(doc, title, draft.meta.doc_id, "standard",
                 draft.meta.version, draft.meta.classification, rtl=rtl)
    _setup_header_footer(doc, draft.meta.doc_id, title,
                         draft.meta.classification, rtl=rtl)

    # Notice page
    _section_heading(doc, "Ownership Notice", rtl=rtl)
    _body_para(doc, draft.notice_en or draft.notice_ar, rtl=rtl)
    doc.add_page_break()

    # Document info
    _section_heading(doc, "Document Information", rtl=rtl)
    info_rows = [
        ("Document Title", title),
        ("Document ID", draft.meta.doc_id),
        ("Document Type", "Cybersecurity Standard"),
        ("Version", draft.meta.version),
        ("Owner", draft.meta.owner),
        ("Classification", draft.meta.classification),
    ]
    _info_table(doc, info_rows, rtl=rtl)

    # Approval summary
    _section_heading(doc, "Document Approval", rtl=rtl)
    _approval_table(doc, draft.approval_stages, rtl=rtl)
    doc.add_page_break()

    # ToC
    _section_heading(doc, "Table of Contents", rtl=rtl)
    for item in ["1.  Abbreviations and Definitions", "2.  Objective",
                 "3.  Scope and Applicability", "4.  Standards",
                 "5.  Exceptions", "6.  Roles and Responsibilities",
                 "7.  Update and Review", "8.  Compliance"]:
        _bullet_para(doc, item, rtl=rtl)
    doc.add_page_break()

    # Section 1
    _section_heading(doc, "1.  Abbreviations and Definitions", rtl=rtl)
    _definitions_table(doc, draft.definitions, rtl=rtl)

    # Section 2
    _section_heading(doc, "2.  Objective", rtl=rtl)
    _body_para(doc, draft.objective_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 3: Scope
    _section_heading(doc, "3.  Scope and Applicability", rtl=rtl)
    _body_para(doc, draft.scope_intro_ar, rtl=rtl)
    for cat_name, items in [
        ("Systems and Applications", draft.scope_systems_ar),
        ("Roles and Permissions",    draft.scope_roles_ar),
        ("Critical Processes",        draft.scope_processes_ar),
        ("Covered Persons and Entities", draft.scope_persons_ar),
    ]:
        if items:
            _section_heading(doc, cat_name, level=2, rtl=rtl)
            for item in items:
                _bullet_para(doc, item, rtl=rtl)
    doc.add_paragraph()

    # Section 4: Standards (dominant)
    _section_heading(doc, "4.  Standards", rtl=rtl)
    for cluster in draft.domain_clusters:
        _section_heading(doc, f"{cluster.cluster_id}  {cluster.title_ar}", level=2, rtl=rtl)

        # Objective + risks info box
        info_p = doc.add_paragraph()
        if rtl:
            _set_rtl(info_p)
        _set_para_spacing(info_p, 0, 3)
        _add_run(info_p, "Objective:  ", bold=True, size_pt=10, color_hex=C_GREEN, font="Calibri")
        _add_run(info_p, cluster.objective_ar, size_pt=10, font="Calibri")

        risk_p = doc.add_paragraph()
        if rtl:
            _set_rtl(risk_p)
        _set_para_spacing(risk_p, 0, 6)
        _add_run(risk_p, "Potential Risks:  ", bold=True, size_pt=10, color_hex="8B0000", font="Calibri")
        _add_run(risk_p, cluster.potential_risks_ar, size_pt=10, font="Calibri")

        for clause in cluster.clauses:
            p = doc.add_paragraph()
            if rtl:
                _set_rtl(p)
            else:
                _para_align(p, "both")
            _set_para_spacing(p, 2, 4)
            pPr = p._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:left"), "360")
            ind.set(qn("w:hanging"), "360")
            pPr.append(ind)
            id_run = p.add_run(f"{clause.clause_id}  ")
            id_run.bold = True
            id_run.font.size = Pt(10)
            id_run.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
            id_run.font.name = "Calibri"
            txt_run = p.add_run(clause.text_ar)
            txt_run.font.size = Pt(11)
            txt_run.font.name = "Calibri"
            if clause.guidance_ar:
                gp = doc.add_paragraph()
                if rtl:
                    _set_rtl(gp)
                _set_para_spacing(gp, 0, 3)
                pPr2 = gp._p.get_or_add_pPr()
                ind2 = OxmlElement("w:ind")
                ind2.set(qn("w:left"), "720")
                pPr2.append(ind2)
                _add_run(gp, "Guidance:  ", bold=True, size_pt=9,
                         color_hex=C_GREEN_LIGHT, font="Calibri")
                _add_run(gp, clause.guidance_ar, italic=True, size_pt=9, font="Calibri")
        doc.add_paragraph()

    # Section 5: Exceptions
    _section_heading(doc, "5.  Exceptions", rtl=rtl)
    _body_para(doc, draft.exceptions_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 6: Roles
    _section_heading(doc, "6.  Roles and Responsibilities", rtl=rtl)
    if draft.roles_responsibilities:
        roles_table = doc.add_table(rows=1 + len(draft.roles_responsibilities), cols=2)
        roles_table.style = "Table Grid"
        _set_thin_borders(roles_table)
        _set_table_width_full(roles_table)
        _make_header_row(roles_table, ["Role", "Responsibilities"], rtl=rtl)
        for i, (role, duties) in enumerate(draft.roles_responsibilities.items()):
            _fill_data_row(roles_table.rows[i + 1], [role, duties], alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()

    # Section 7 + 8
    _section_heading(doc, "7.  Update and Review", rtl=rtl)
    _body_para(doc, draft.update_review_ar, rtl=rtl)
    doc.add_paragraph()

    _section_heading(doc, "8.  Compliance with the Standard", rtl=rtl)
    _body_para(doc, draft.compliance_ar, rtl=rtl)
    doc.add_paragraph()

    _section_heading(doc, "Version Control", rtl=rtl)
    _version_table(doc, draft.version_rows, rtl=rtl)

    suffix = "_en" if lang == "en" else "_ar"
    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}{suffix}.docx")
    doc.save(out_path)
    return out_path


# ═══════════════════════════════════════════════════════════════════════════════
# PROCEDURE — Professional Renderer
# ═══════════════════════════════════════════════════════════════════════════════

def render_pro_procedure(draft: MinistryProcedureDraft, output_dir: str,
                         lang: str = "en") -> str:
    os.makedirs(output_dir, exist_ok=True)
    rtl = (lang == "ar")
    doc = _new_doc(rtl=rtl)
    title = draft.meta.title_en if lang == "en" else draft.meta.title_ar

    _build_cover(doc, title, draft.meta.doc_id, "procedure",
                 draft.meta.version, draft.meta.classification, rtl=rtl)
    _setup_header_footer(doc, draft.meta.doc_id, title,
                         draft.meta.classification, rtl=rtl)

    # Document info
    _section_heading(doc, "Document Information", rtl=rtl)
    info_rows = [
        ("Document Title",    title),
        ("Document ID",       draft.meta.doc_id),
        ("Document Type",     "Cybersecurity Procedure"),
        ("Parent Policy",     draft.parent_policy_id or "—"),
        ("Parent Standard",   draft.parent_standard_id or "—"),
        ("Version",           draft.meta.version),
        ("Owner",             draft.meta.owner),
        ("Classification",    draft.meta.classification),
    ]
    _info_table(doc, info_rows, rtl=rtl)

    # Approval
    _section_heading(doc, "Document Approval", rtl=rtl)
    _approval_table(doc, draft.approval_stages, rtl=rtl)
    doc.add_page_break()

    # ToC
    _section_heading(doc, "Table of Contents", rtl=rtl)
    for item in [
        "1.  Definitions and Abbreviations", "2.  Procedure Objective",
        "3.  Scope and Applicability", "4.  Roles and Responsibilities",
        "5.  Procedure Overview", "6.  Triggers, Prerequisites, Inputs and Tools",
        "7.  Detailed Procedure Steps", "8.  Decision Points, Exceptions and Escalations",
        "9.  Outputs, Records, Evidence and Forms",
        "10. Time Controls, SLAs and Control Checkpoints",
        "11. Related Documents", "12. Effective Date",
        "13. Procedure Review", "14. Review and Approval", "15. Version Control",
    ]:
        _bullet_para(doc, item, rtl=rtl)
    doc.add_page_break()

    # Section 1
    _section_heading(doc, "1.  Definitions and Abbreviations", rtl=rtl)
    _definitions_table(doc, draft.definitions, rtl=rtl)

    # Section 2
    _section_heading(doc, "2.  Procedure Objective", rtl=rtl)
    _body_para(doc, draft.objective_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 3
    _section_heading(doc, "3.  Scope and Applicability", rtl=rtl)
    _body_para(doc, draft.scope_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 4: Roles
    _section_heading(doc, "4.  Roles and Responsibilities", rtl=rtl)
    for role_item in draft.roles:
        p = doc.add_paragraph()
        if rtl:
            _set_rtl(p)
        _set_para_spacing(p, 6, 2)
        _add_run(p, role_item.role_ar, bold=True, size_pt=11,
                 color_hex=C_GREEN, font="Calibri")
        for resp in role_item.responsibilities_ar:
            _bullet_para(doc, resp, rtl=rtl)
    doc.add_paragraph()

    # Section 5
    _section_heading(doc, "5.  Procedure Overview", rtl=rtl)
    _body_para(doc, draft.overview_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 6
    _section_heading(doc, "6.  Triggers, Prerequisites, Inputs and Tools", rtl=rtl)
    for label, items in [
        ("Triggers",          draft.triggers_ar),
        ("Prerequisites",     draft.prerequisites_ar),
        ("Inputs",            draft.inputs_ar),
        ("Tools and Systems", draft.tools_ar),
    ]:
        if items:
            _section_heading(doc, label, level=2, rtl=rtl)
            for item in items:
                _bullet_para(doc, item, rtl=rtl)
    doc.add_paragraph()

    # Section 7: Detailed Steps (dominant)
    _section_heading(doc, "7.  Detailed Procedure Steps", rtl=rtl)
    for phase in draft.phases:
        _section_heading(doc, f"{phase.phase_id}  {phase.phase_title_ar}", level=2, rtl=rtl)
        if phase.phase_objective_ar:
            _body_para(doc, phase.phase_objective_ar, rtl=rtl)

        if phase.steps:
            # Steps table with 6 columns
            cols = ["Step ID", "Actor", "Action", "System / Tool", "Output", "Evidence"]
            steps_table = doc.add_table(rows=1 + len(phase.steps), cols=6)
            steps_table.style = "Table Grid"
            _set_thin_borders(steps_table)
            _set_table_width_full(steps_table)
            _make_header_row(steps_table, cols, rtl=rtl)
            for i, step in enumerate(phase.steps):
                row_texts = [step.step_id, step.actor_ar, step.action_ar,
                             step.system_ar, step.output_ar, step.evidence_ar]
                _fill_data_row(steps_table.rows[i + 1], row_texts,
                               alt=(i % 2 == 0), rtl=rtl)
        doc.add_paragraph()

    # Section 8
    _section_heading(doc, "8.  Decision Points, Exceptions and Escalations", rtl=rtl)
    if draft.decision_points_ar:
        _section_heading(doc, "Decision Points", level=2, rtl=rtl)
        for dp in draft.decision_points_ar:
            _bullet_para(doc, dp, rtl=rtl)
    _section_heading(doc, "Exceptions", level=2, rtl=rtl)
    _body_para(doc, draft.exceptions_ar, rtl=rtl)
    _section_heading(doc, "Escalation Path", level=2, rtl=rtl)
    _body_para(doc, draft.escalation_ar, rtl=rtl)
    doc.add_paragraph()

    # Section 9
    _section_heading(doc, "9.  Outputs, Records, Evidence and Forms", rtl=rtl)
    for label, items in [
        ("Outputs",  draft.outputs_ar),
        ("Records",  draft.records_ar),
        ("Evidence", draft.evidence_ar),
        ("Forms",    draft.forms_ar),
    ]:
        if items:
            _section_heading(doc, label, level=2, rtl=rtl)
            for item in items:
                _bullet_para(doc, item, rtl=rtl)
    doc.add_paragraph()

    # Section 10
    _section_heading(doc, "10. Time Controls, SLAs and Control Checkpoints", rtl=rtl)
    for tc in draft.time_controls_ar:
        _bullet_para(doc, tc, rtl=rtl)
    doc.add_paragraph()

    # Section 11
    _section_heading(doc, "11. Related Documents", rtl=rtl)
    if draft.related_docs:
        rel_table = doc.add_table(rows=1 + len(draft.related_docs), cols=2)
        rel_table.style = "Table Grid"
        _set_thin_borders(rel_table)
        _set_table_width_full(rel_table)
        _make_header_row(rel_table, ["Document Title", "Type"], rtl=rtl)
        for i, rd in enumerate(draft.related_docs):
            _fill_data_row(rel_table.rows[i + 1], [rd.title_ar, rd.ref_type],
                           alt=(i % 2 == 0), rtl=rtl)
    doc.add_paragraph()

    # Sections 12-15
    _section_heading(doc, "12. Effective Date", rtl=rtl)
    _body_para(doc, draft.effective_date_ar, rtl=rtl)
    doc.add_paragraph()

    _section_heading(doc, "13. Procedure Review", rtl=rtl)
    _body_para(doc, draft.review_ar, rtl=rtl)
    doc.add_paragraph()

    _section_heading(doc, "14. Review and Approval", rtl=rtl)
    _approval_table(doc, draft.approval_stages, rtl=rtl)

    _section_heading(doc, "15. Version Control", rtl=rtl)
    _version_table(doc, draft.version_rows, rtl=rtl)

    suffix = "_en" if lang == "en" else "_ar"
    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}{suffix}.docx")
    doc.save(out_path)
    return out_path
