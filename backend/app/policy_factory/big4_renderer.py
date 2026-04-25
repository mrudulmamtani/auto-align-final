"""
BIG4-Quality DOCX Renderer
----------------------------
Strict color codes (from policy.json / standard.json / procedure.json):
  Primary Dark Green : #0F4D43
  Secondary Green    : #2E7D6B
  Neutral Light      : #F4F4F4
  Border Gray        : #D9D9D9
  White              : #FFFFFF
  Text Dark          : #1A1A1A

Produces output files: {doc_id}_big4.docx
"""
from __future__ import annotations

import os
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from .langchain_drafter import (
    PolicySpec, StandardSpec, ProcedureSpec,
    TermDefinition, ControlTrace, PolicyElement,
    DomainBlock, Requirement, ProcedurePhase, ProcedureStep, RoleRow,
)

# ── Brand palette ─────────────────────────────────────────────────────────────
C_PRIMARY   = "0F4D43"   # dark green  — headings, cover bg, table headers
C_SECONDARY = "2E7D6B"   # mid green   — sub-headings, accents
C_NEUTRAL   = "F4F4F4"   # light gray  — label cells, alt rows
C_BORDER    = "D9D9D9"   # border gray — table borders
C_WHITE     = "FFFFFF"
C_TEXT      = "1A1A1A"
C_MUTED     = "5A5A5A"   # footer / watermark text

ORG_NAME  = "Ministry of Communications and Information Technology"
ORG_SHORT = "MCIT"
FONT_NAME = "Calibri"


# ══════════════════════════════════════════════════════════════════════════════
# Low-level OOXML helpers
# ══════════════════════════════════════════════════════════════════════════════

def _rgb(hex6: str) -> RGBColor:
    return RGBColor(int(hex6[:2], 16), int(hex6[2:4], 16), int(hex6[4:], 16))


def _cell_bg(cell, hex6: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex6)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _row_height(row, cm: float) -> None:
    trPr = row._tr.get_or_add_trPr()
    for old in trPr.findall(qn("w:trHeight")):
        trPr.remove(old)
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(cm * 567)))
    h.set(qn("w:hRule"), "exact")
    trPr.append(h)


def _no_borders(table) -> None:
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "none"); b.set(qn("w:sz"), "0")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), "auto")
        tblB.append(b)
    tblPr.append(tblB)


def _thin_borders(table) -> None:
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblBorders")):
        tblPr.remove(old)
    tblB = OxmlElement("w:tblBorders")
    for side in ("top","left","bottom","right","insideH","insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0"); b.set(qn("w:color"), C_BORDER)
        tblB.append(b)
    tblPr.append(tblB)


def _table_full_width(table) -> None:
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl.insert(0, tblPr)
    for old in tblPr.findall(qn("w:tblW")):
        tblPr.remove(old)
    w = OxmlElement("w:tblW")
    w.set(qn("w:w"), "9898"); w.set(qn("w:type"), "pct")
    tblPr.append(w)


def _para_spacing(para, before: int = 0, after: int = 6,
                  line: int = 276) -> None:
    pPr = para._p.get_or_add_pPr()
    sp  = pPr.find(qn("w:spacing"))
    if sp is None:
        sp = OxmlElement("w:spacing"); pPr.append(sp)
    sp.set(qn("w:before"), str(before * 20))
    sp.set(qn("w:after"),  str(after  * 20))
    sp.set(qn("w:line"),   str(line))
    sp.set(qn("w:lineRule"), "auto")


def _para_align(para, val: str) -> None:
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:jc")):
        pPr.remove(old)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), val)
    pPr.append(jc)


def _run(para, text: str, bold=False, italic=False, size_pt=11,
         color=C_TEXT, font=FONT_NAME) -> None:
    r = para.add_run(text)
    r.bold   = bold
    r.italic = italic
    r.font.name = font
    r.font.size = Pt(size_pt)
    r.font.color.rgb = _rgb(color)


def _indent(para, left_cm: float, hanging_cm: float = 0) -> None:
    pPr = para._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:ind")):
        pPr.remove(old)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"),    str(int(left_cm * 567)))
    if hanging_cm:
        ind.set(qn("w:hanging"), str(int(hanging_cm * 567)))
    pPr.append(ind)


def _para_bottom_border(para, color=C_PRIMARY, sz=8) -> None:
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"),  str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _field_code(para, instr: str) -> None:
    r = para.add_run()
    for tag, text in [("begin", None), ("instrText", instr), ("end", None)]:
        if tag == "instrText":
            el = OxmlElement("w:instrText")
            el.set(qn("xml:space"), "preserve"); el.text = text
        else:
            el = OxmlElement("w:fldChar"); el.set(qn("w:fldCharType"), tag)
        r._r.append(el)


# ── DRAFT watermark (VML) ─────────────────────────────────────────────────────
_WM_VML = """<v:shape xmlns:v="urn:schemas-microsoft-com:vml"
  xmlns:o="urn:schemas-microsoft-com:office:office"
  xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"
  id="WatermarkObj" o:spid="_x0000_s2052" type="#_x0000_t136"
  style="position:absolute;margin-left:0;margin-top:0;width:527.85pt;height:131.95pt;
         z-index:-251655168;mso-position-horizontal:center;
         mso-position-horizontal-relative:margin;mso-position-vertical:center;
         mso-position-vertical-relative:margin"
  fillcolor="#CCCCCC" stroked="f">
  <v:textpath style='font-family:"Calibri";font-size:1pt;font-weight:bold'
              string="DRAFT" trim="t" on="t"/>
</v:shape>"""

def _add_watermark(header) -> None:
    if not header.paragraphs:
        header.add_paragraph()
    para = header.paragraphs[0]
    r    = OxmlElement("w:r")
    rPr  = OxmlElement("w:rPr")
    rPr.append(OxmlElement("w:noProof"))
    r.append(rPr)
    pict = OxmlElement("w:pict")
    pict.append(etree.fromstring(_WM_VML))
    r.append(pict)
    para._p.append(r)


# ══════════════════════════════════════════════════════════════════════════════
# Document factory
# ══════════════════════════════════════════════════════════════════════════════

def _new_doc() -> Document:
    doc = Document()
    sty = doc.styles["Normal"]
    sty.font.name = FONT_NAME
    sty.font.size = Pt(11)
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(2.54)
        sec.right_margin  = Cm(2.54)
    return doc


# ══════════════════════════════════════════════════════════════════════════════
# Cover page  (full-page dark green, title center-lower third)
# ══════════════════════════════════════════════════════════════════════════════

def _build_cover(doc: Document, title: str, doc_id: str,
                 doc_type: str, version: str, classification: str) -> None:
    """3-band borderless cover: green top | green mid with title | green footer."""
    tbl = doc.add_table(rows=3, cols=1)
    _no_borders(tbl)
    _table_full_width(tbl)

    # ── Row 1 — top green band (org name + type tag) ──
    r0 = tbl.rows[0]
    _row_height(r0, 9.0)
    _cell_bg(r0.cells[0], C_PRIMARY)
    c0 = r0.cells[0]
    c0.paragraphs[0].clear()

    sp = c0.add_paragraph(); _para_spacing(sp, 30, 0)

    org_p = c0.add_paragraph()
    _para_align(org_p, "center"); _para_spacing(org_p, 0, 6)
    _run(org_p, ORG_NAME, bold=True, size_pt=13, color=C_WHITE)

    tag_label = {
        "policy":    "CYBERSECURITY POLICY",
        "standard":  "CYBERSECURITY STANDARD",
        "procedure": "CYBERSECURITY PROCEDURE",
    }.get(doc_type, "DOCUMENT")
    tag_p = c0.add_paragraph()
    _para_align(tag_p, "center"); _para_spacing(tag_p, 4, 4)
    _run(tag_p, tag_label, size_pt=10, color="A8D8CF")

    # decorative rule
    rule_p = c0.add_paragraph()
    _para_align(rule_p, "center"); _para_spacing(rule_p, 10, 0)
    _run(rule_p, "─" * 46, size_pt=9, color="4DA88E")

    # ── Row 2 — white title area (center-lower third per schema) ──
    r1 = tbl.rows[1]
    _row_height(r1, 13.0)
    _cell_bg(r1.cells[0], C_WHITE)
    c1 = r1.cells[0]
    c1.paragraphs[0].clear()

    # Push title to lower third with spacer
    c1.add_paragraph(); c1.add_paragraph(); c1.add_paragraph()
    id_p = c1.add_paragraph()
    _para_align(id_p, "center"); _para_spacing(id_p, 0, 4)
    _run(id_p, doc_id, size_pt=13, color=C_SECONDARY)

    title_p = c1.add_paragraph()
    _para_align(title_p, "center"); _para_spacing(title_p, 6, 8)
    _run(title_p, title, bold=True, size_pt=22, color=C_PRIMARY)

    # decorative green rule under title
    rule2 = c1.add_paragraph()
    _para_align(rule2, "center"); _para_spacing(rule2, 4, 4)
    _para_bottom_border(rule2, color=C_PRIMARY, sz=12)

    # ── Row 3 — dark green metadata footer ──
    r2 = tbl.rows[2]
    _row_height(r2, 5.5)
    _cell_bg(r2.cells[0], C_PRIMARY)
    c2 = r2.cells[0]
    c2.paragraphs[0].clear()

    sp2 = c2.add_paragraph(); _para_spacing(sp2, 14, 0)
    meta_p = c2.add_paragraph()
    _para_align(meta_p, "center"); _para_spacing(meta_p, 0, 4)
    _run(meta_p,
         f"Version {version}     |     Issue Date: [Date of Approval]     |     {classification}",
         size_pt=9, color=C_WHITE)

    doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# Header / Footer (internal pages)
# ══════════════════════════════════════════════════════════════════════════════

def _setup_hf(doc: Document, doc_id: str,
              classification: str) -> None:
    for sec in doc.sections:
        sec.different_first_page_header_footer = True

        # First page (cover): blank
        sec.first_page_header.paragraphs[0].clear()
        sec.first_page_footer.paragraphs[0].clear()

        # ── Default header ──
        hdr = sec.header
        hdr.is_linked_to_previous = False
        if hdr.paragraphs:
            hdr.paragraphs[0].clear()
        else:
            hdr.add_paragraph()

        _add_watermark(hdr)

        # 2-col table: left = org short | right = doc ref
        ht = hdr.add_table(rows=1, cols=2, width=Inches(6.3))
        _no_borders(ht)
        _table_full_width(ht)

        lp = ht.rows[0].cells[0].paragraphs[0]
        lp.clear(); _para_spacing(lp, 0, 0)
        _run(lp, f"{ORG_SHORT} | Cybersecurity",
             bold=True, size_pt=8, color=C_PRIMARY)

        rp = ht.rows[0].cells[1].paragraphs[0]
        rp.clear(); _para_align(rp, "right"); _para_spacing(rp, 0, 0)
        _run(rp, f"{doc_id}  |  v1.0  |  {classification}",
             size_pt=8, color=C_SECONDARY)

        # Green bottom border on header paragraph
        hp = hdr.paragraphs[0]
        _para_bottom_border(hp, color=C_PRIMARY, sz=8)

        # ── Default footer ──
        ftr = sec.footer
        ftr.is_linked_to_previous = False
        if ftr.paragraphs:
            ftr.paragraphs[0].clear()
        else:
            ftr.add_paragraph()

        ft = ftr.add_table(rows=1, cols=3, width=Inches(6.3))
        _no_borders(ft)
        _table_full_width(ft)

        # Left: classification
        fl = ft.rows[0].cells[0].paragraphs[0]
        fl.clear(); _para_spacing(fl, 0, 0)
        _run(fl, classification, bold=True, size_pt=8, color=C_PRIMARY)

        # Center: Page X of Y
        fc = ft.rows[0].cells[1].paragraphs[0]
        fc.clear(); _para_align(fc, "center"); _para_spacing(fc, 0, 0)
        _run(fc, "Page ", size_pt=8, color=C_TEXT)
        _field_code(fc, " PAGE ")
        _run(fc, " of ", size_pt=8, color=C_TEXT)
        _field_code(fc, " NUMPAGES ")

        # Right: doc_id
        fr = ft.rows[0].cells[2].paragraphs[0]
        fr.clear(); _para_align(fr, "right"); _para_spacing(fr, 0, 0)
        _run(fr, doc_id, size_pt=8, color=C_SECONDARY)

        # Green top border on footer paragraph
        fp0 = ftr.paragraphs[0]
        pPr = fp0._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        top  = OxmlElement("w:top")
        top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "8")
        top.set(qn("w:space"), "1");    top.set(qn("w:color"), C_PRIMARY)
        pBdr.append(top); pPr.append(pBdr)


# ══════════════════════════════════════════════════════════════════════════════
# Content helpers
# ══════════════════════════════════════════════════════════════════════════════

def _h1(doc: Document, number: str, text: str) -> None:
    """Numbered H1 with green bottom rule."""
    p = doc.add_paragraph()
    _para_spacing(p, before=18, after=4)
    _run(p, f"{number}  {text}", bold=True, size_pt=13, color=C_PRIMARY)
    _para_bottom_border(p, color=C_SECONDARY, sz=6)


def _h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, before=10, after=2)
    _run(p, text, bold=True, size_pt=11, color=C_SECONDARY)


def _body(doc: Document, text: str, indent_cm: float = 0) -> None:
    p = doc.add_paragraph()
    _para_align(p, "both")
    _para_spacing(p, 0, 6)
    if indent_cm:
        _indent(p, indent_cm)
    _run(p, text, size_pt=11)


def _bullet(doc: Document, text: str, indent_cm: float = 0.6) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, 0, 3)
    _indent(p, indent_cm + 0.3, hanging_cm=0.3)
    _run(p, f"•  {text}", size_pt=10.5)


def _numbered_item(doc: Document, number: str, text: str,
                   indent_cm: float = 0.5) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, 2, 4)
    _para_align(p, "both")
    _indent(p, indent_cm + 0.5, hanging_cm=0.5)
    _run(p, f"{number}  ", bold=True, size_pt=11, color=C_PRIMARY)
    _run(p, text, size_pt=11)


def _spacer(doc: Document) -> None:
    p = doc.add_paragraph()
    _para_spacing(p, 0, 0)


# ── Table builders ────────────────────────────────────────────────────────────

def _make_table(doc: Document, headers: list[str]) -> "Table":
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    _thin_borders(tbl)
    _table_full_width(tbl)
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        _cell_bg(c, C_PRIMARY)
        c.paragraphs[0].clear()
        p = c.paragraphs[0]
        _para_align(p, "center")
        _para_spacing(p, 3, 3)
        _run(p, h, bold=True, size_pt=9, color=C_WHITE)
    return tbl


def _add_row(tbl, texts: list[str], alt: bool = False) -> None:
    row = tbl.add_row()
    bg  = C_NEUTRAL if alt else C_WHITE
    for i, cell in enumerate(row.cells):
        _cell_bg(cell, bg)
        cell.paragraphs[0].clear()
        p = cell.paragraphs[0]
        _para_spacing(p, 2, 2)
        _run(p, texts[i] if i < len(texts) else "", size_pt=9)


def _definitions_table(doc: Document, defs: list[TermDefinition]) -> None:
    if not defs:
        return
    tbl = _make_table(doc, ["Term", "Definition"])
    for i, d in enumerate(defs):
        _add_row(tbl, [d.term, d.definition], alt=(i % 2 == 0))
    _spacer(doc)


def _approval_table(doc: Document) -> None:
    stages = [
        ("Prepared by",        "Director, Cybersecurity Department", "[Name]", "[Date]"),
        ("Reviewed by",        "Deputy Director, Cybersecurity Dept", "[Name]", "[Date]"),
        ("Recommended by",     "VP, Technology Governance",          "[Name]", "[Date]"),
        ("Approved by",        "Minister",                           "[Name]", "[Date]"),
    ]
    tbl = _make_table(doc, ["Stage", "Role / Title", "Name", "Date"])
    for i, (st, role, name, date) in enumerate(stages):
        _add_row(tbl, [st, role, name, date], alt=(i % 2 == 0))
    _spacer(doc)


def _version_table(doc: Document, version: str = "1.0") -> None:
    tbl = _make_table(doc,
        ["Version", "Update Type", "Summary", "Updated By", "Approval Date"])
    _add_row(tbl, [version, "Initial Release",
                   "First approved version.",
                   "Cybersecurity Department", "[Date]"])
    _spacer(doc)


def _info_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"
    _thin_borders(tbl)
    _table_full_width(tbl)
    for i, (label, value) in enumerate(rows):
        lc = tbl.rows[i].cells[0]
        vc = tbl.rows[i].cells[1]
        _cell_bg(lc, C_NEUTRAL)
        lc.paragraphs[0].clear()
        lp = lc.paragraphs[0]
        _para_spacing(lp, 2, 2)
        _run(lp, label, bold=True, size_pt=9, color=C_TEXT)
        vc.paragraphs[0].clear()
        vp = vc.paragraphs[0]
        _para_spacing(vp, 2, 2)
        _run(vp, value, size_pt=9)
    _spacer(doc)


def _control_mapping_table(doc: Document, elements: list) -> None:
    """Control Traceability Annex table."""
    _h1(doc, "Annex A", "Control Traceability Matrix")
    tbl = _make_table(doc, ["Clause / Req", "Statement", "Framework", "Control ID", "Source Reference"])
    for el in elements:
        stmt  = ""
        trace = []
        num   = ""
        if hasattr(el, "statement"):   # PolicyElement / Requirement
            stmt  = el.statement
            trace = el.trace if hasattr(el, "trace") else []
            num   = getattr(el, "element_no", getattr(el, "req_id", ""))
        for t in trace:
            _add_row(tbl, [num, stmt[:120], t.framework, t.control_id, t.source_ref])
    _spacer(doc)


# ══════════════════════════════════════════════════════════════════════════════
# POLICY renderer
# ══════════════════════════════════════════════════════════════════════════════

def render_policy(draft: PolicySpec, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc()

    _build_cover(doc, draft.title, draft.doc_id, "policy",
                 draft.version, draft.classification)
    _setup_hf(doc, draft.doc_id, draft.classification)

    # ── Document Information ──
    _h1(doc, "Doc. Info", "Document Information")
    _info_table(doc, [
        ("Document Title",      draft.title),
        ("Document ID",         draft.doc_id),
        ("Document Type",       "Cybersecurity Policy"),
        ("Version",             draft.version),
        ("Owner",               draft.owner),
        ("Classification",      draft.classification),
        ("Issue Date",          "[Date of Approval]"),
        ("Effective Date",      "[Date of Approval]"),
        ("Review Due",          "Annually or upon material change"),
    ])

    # ── TOC ──
    _h1(doc, "ToC", "Table of Contents")
    for n, s in enumerate(["Objectives", "Scope and Applicability",
                            "Policy Elements", "Roles and Responsibilities",
                            "Policy Compliance", "Exceptions"], 1):
        p = doc.add_paragraph()
        _para_spacing(p, 1, 1)
        _indent(p, 0.3)
        _run(p, f"{n}.  {s}", size_pt=11, color=C_SECONDARY)
    doc.add_page_break()

    # ── Section 1: Definitions ──
    _h1(doc, "1.", "Definitions and Abbreviations")
    _definitions_table(doc, draft.definitions)

    # ── Section 2: Objectives ──
    _h1(doc, "2.", "Objectives")
    _body(doc, draft.objectives)
    _spacer(doc)

    # ── Section 3: Scope ──
    _h1(doc, "3.", "Scope and Applicability")
    _body(doc, draft.scope)
    _spacer(doc)

    # ── Section 4: Policy Elements (dominant section) ──
    _h1(doc, "4.", "Policy Elements")
    _h2(doc, "General Provisions")
    for el in draft.policy_elements:
        # Element number + statement
        p = doc.add_paragraph()
        _para_spacing(p, 4, 3)
        _para_align(p, "both")
        _indent(p, 0.6, hanging_cm=0.6)
        _run(p, f"{el.element_no}.  ", bold=True, size_pt=11, color=C_PRIMARY)
        _run(p, el.statement, size_pt=11)

        # Sub-items (element 2 policy catalogue)
        if el.sub_items:
            for sub in el.sub_items:
                _bullet(doc, sub, indent_cm=0.9)

        # Traceability tags
        if el.trace:
            ids_str = "  ".join(f"[{t.control_id}]" for t in el.trace)
            tag_p = doc.add_paragraph()
            _para_spacing(tag_p, 0, 4)
            _indent(tag_p, 1.2)
            _run(tag_p, f"NCA Control Reference: {ids_str}",
                 italic=True, size_pt=8.5, color=C_SECONDARY)
    _spacer(doc)

    # ── Section 5: Roles ──
    _h1(doc, "5.", "Roles and Responsibilities")
    _body(doc, draft.roles_responsibilities)
    _spacer(doc)

    # ── Section 6: Policy Compliance ──
    _h1(doc, "6.", "Policy Compliance")
    for i, clause in enumerate(draft.compliance_clauses, 1):
        _numbered_item(doc, f"{i}.", clause)
    _spacer(doc)

    # ── Section 7: Exceptions ──
    _h1(doc, "7.", "Exceptions")
    _body(doc, draft.exceptions)
    _spacer(doc)

    # ── Section 8: Effective Date & Review ──
    _h1(doc, "8.", "Effective Date and Review")
    _body(doc, "This policy comes into force on the date of its formal approval "
               "and remains in effect until superseded or withdrawn.")
    _body(doc, "This policy shall be reviewed annually, or upon material changes "
               "in regulatory requirements, technology landscape, or organisational "
               "structure. All amendments require formal approval.")
    _spacer(doc)

    # ── Section 9: Review and Approval ──
    _h1(doc, "9.", "Review and Approval")
    _approval_table(doc)

    # ── Section 10: Version Control ──
    _h1(doc, "10.", "Version Control")
    _version_table(doc, draft.version)

    # ── Closing note ──
    _spacer(doc)
    note_p = doc.add_paragraph()
    _para_align(note_p, "center")
    _para_spacing(note_p, 8, 8)
    _run(note_p, draft.closing_note, bold=True, size_pt=10, color=C_PRIMARY)

    # ── Annex A: Control Traceability ──
    doc.add_page_break()
    _control_mapping_table(doc, draft.policy_elements)

    path = os.path.join(output_dir, f"{draft.doc_id}_big4.docx")
    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════════════════════
# STANDARD renderer
# ══════════════════════════════════════════════════════════════════════════════

def render_standard(draft: StandardSpec, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc()

    _build_cover(doc, draft.title, draft.doc_id, "standard",
                 draft.version, draft.classification)
    _setup_hf(doc, draft.doc_id, draft.classification)

    # ── Ownership Notice ──
    _h1(doc, "", "Ownership Notice")
    _body(doc, draft.notice)
    doc.add_page_break()

    # ── Document Info + Approval ──
    _h1(doc, "", "Document Information")
    _info_table(doc, [
        ("Document Title",   draft.title),
        ("Document ID",      draft.doc_id),
        ("Document Type",    "Cybersecurity Standard"),
        ("Version",          draft.version),
        ("Owner",            draft.owner),
        ("Classification",   draft.classification),
        ("Issue Date",       "[Date of Approval]"),
    ])
    _h2(doc, "Document Approval")
    _approval_table(doc)
    _h2(doc, "Revision History")
    _version_table(doc, draft.version)
    doc.add_page_break()

    # ── TOC ──
    _h1(doc, "ToC", "Table of Contents")
    for n, s in enumerate(["Abbreviations and Definitions", "Objective",
                            "Scope of Work and Applicability", "Standards",
                            "Exceptions", "Roles and Responsibilities",
                            "Update and Review",
                            "Compliance with the Standard"], 1):
        p = doc.add_paragraph()
        _para_spacing(p, 1, 1)
        _indent(p, 0.3)
        _run(p, f"{n}.  {s}", size_pt=11, color=C_SECONDARY)
    doc.add_page_break()

    # ── Section 1: Definitions ──
    _h1(doc, "1.", "Abbreviations and Definitions")
    _definitions_table(doc, draft.definitions)

    # ── Section 2: Objective ──
    _h1(doc, "2.", "Objective")
    _body(doc, draft.objectives)
    _spacer(doc)

    # ── Section 3: Scope ──
    _h1(doc, "3.", "Scope of Work and Applicability")
    _body(doc, draft.scope)
    _spacer(doc)

    # ── Section 4: Standards (dominant) ──
    _h1(doc, "4.", "Standards")
    for domain in draft.domains:
        _h2(doc, f"4.{domain.domain_no}  {domain.title}")

        obj_p = doc.add_paragraph()
        _para_spacing(obj_p, 2, 2)
        _para_align(obj_p, "both")
        _run(obj_p, "Objective:  ", bold=True, size_pt=10.5, color=C_PRIMARY)
        _run(obj_p, domain.objective, size_pt=10.5)

        risk_p = doc.add_paragraph()
        _para_spacing(risk_p, 0, 6)
        _para_align(risk_p, "both")
        _run(risk_p, "Potential Risks:  ", bold=True, size_pt=10.5, color="8B0000")
        _run(risk_p, domain.potential_risks, size_pt=10.5)

        _h2(doc, "Requirements")
        for req in domain.requirements:
            p = doc.add_paragraph()
            _para_spacing(p, 3, 2)
            _para_align(p, "both")
            _indent(p, 0.6, hanging_cm=0.6)
            _run(p, f"{req.req_id}  ", bold=True, size_pt=11, color=C_PRIMARY)
            _run(p, req.statement, size_pt=11)

            if req.guidance:
                gp = doc.add_paragraph()
                _para_spacing(gp, 0, 3)
                _indent(gp, 1.1)
                _run(gp, "Guidance:  ", bold=True, size_pt=9, color=C_SECONDARY)
                _run(gp, req.guidance, italic=True, size_pt=9)

            if req.trace:
                ids_str = "  ".join(f"[{t.control_id}]" for t in req.trace)
                tag_p = doc.add_paragraph()
                _para_spacing(tag_p, 0, 4)
                _indent(tag_p, 1.1)
                _run(tag_p, f"NCA Reference: {ids_str}",
                     italic=True, size_pt=8.5, color=C_SECONDARY)
        _spacer(doc)

    # ── Section 5: Exceptions ──
    _h1(doc, "5.", "Exceptions")
    _body(doc, draft.exceptions)
    _spacer(doc)

    # ── Section 6: Roles ──
    _h1(doc, "6.", "Roles and Responsibilities")
    for item in draft.roles_responsibilities:
        _numbered_item(doc, "—", item)
    _spacer(doc)

    # ── Section 7: Update & Review ──
    _h1(doc, "7.", "Update and Review")
    _body(doc, draft.update_review)
    _spacer(doc)

    # ── Section 8: Compliance ──
    _h1(doc, "8.", "Compliance with the Standard")
    for i, c in enumerate(draft.compliance_clauses, 1):
        _numbered_item(doc, f"{i}.", c)
    _spacer(doc)

    # ── Closing note ──
    note_p = doc.add_paragraph()
    _para_align(note_p, "center")
    _para_spacing(note_p, 12, 12)
    _run(note_p, draft.closing_note, bold=True, size_pt=10, color=C_PRIMARY)

    # ── Annex A: Control Traceability ──
    doc.add_page_break()
    all_reqs = [req for dom in draft.domains for req in dom.requirements]
    _control_mapping_table(doc, all_reqs)

    path = os.path.join(output_dir, f"{draft.doc_id}_big4.docx")
    doc.save(path)
    return path


# ══════════════════════════════════════════════════════════════════════════════
# PROCEDURE renderer
# ══════════════════════════════════════════════════════════════════════════════

def render_procedure(draft: ProcedureSpec, output_dir: str) -> str:
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc()

    _build_cover(doc, draft.title, draft.doc_id, "procedure",
                 draft.version, draft.classification)
    _setup_hf(doc, draft.doc_id, draft.classification)

    # ── Document Info ──
    _h1(doc, "", "Document Information")
    _info_table(doc, [
        ("Document Title",    draft.title),
        ("Document ID",       draft.doc_id),
        ("Document Type",     "Cybersecurity Procedure"),
        ("Parent Policy",     draft.parent_policy),
        ("Parent Standard",   draft.parent_standard),
        ("Version",           draft.version),
        ("Owner",             draft.owner),
        ("Classification",    draft.classification),
        ("Issue Date",        "[Date of Approval]"),
    ])
    _h2(doc, "Document Approval")
    _approval_table(doc)
    doc.add_page_break()

    # ── TOC ──
    _h1(doc, "ToC", "Table of Contents")
    toc_items = [
        "Definitions and Abbreviations", "Procedure Objective",
        "Scope and Applicability", "Roles and Responsibilities",
        "Procedure Overview", "Triggers, Prerequisites, Inputs and Tools",
        "Detailed Procedure Steps",
        "Decision Points, Exceptions and Escalations",
        "Outputs, Records, Evidence and Forms",
        "Time Controls, SLAs and Control Checkpoints",
        "Related Documents", "Effective Date", "Procedure Review",
        "Review and Approval", "Version Control",
    ]
    for n, s in enumerate(toc_items, 1):
        p = doc.add_paragraph()
        _para_spacing(p, 1, 1); _indent(p, 0.3)
        _run(p, f"{n}.  {s}", size_pt=11, color=C_SECONDARY)
    doc.add_page_break()

    # ── Section 1: Definitions ──
    _h1(doc, "1.", "Definitions and Abbreviations")
    _definitions_table(doc, draft.definitions)

    # ── Section 2: Objective ──
    _h1(doc, "2.", "Procedure Objective")
    _body(doc, draft.objective)
    _spacer(doc)

    # ── Section 3: Scope ──
    _h1(doc, "3.", "Scope and Applicability")
    _body(doc, draft.scope)
    _spacer(doc)

    # ── Section 4: Roles ──
    _h1(doc, "4.", "Roles and Responsibilities")
    tbl = _make_table(doc, ["Role", "Responsibility in this Procedure"])
    for i, r in enumerate(draft.roles):
        _add_row(tbl, [r.role, r.responsibility], alt=(i % 2 == 0))
    _spacer(doc)

    # ── Section 5: Overview ──
    _h1(doc, "5.", "Procedure Overview")
    _body(doc, draft.overview)
    _spacer(doc)

    # ── Section 6: Triggers / Prerequisites / Inputs / Tools ──
    _h1(doc, "6.", "Triggers, Prerequisites, Inputs and Tools")

    for label, items in [
        ("Triggers",      draft.triggers),
        ("Prerequisites", draft.prerequisites),
        ("Inputs",        draft.inputs),
        ("Tools and Systems", draft.tools),
    ]:
        if items:
            _h2(doc, label)
            for item in items:
                _bullet(doc, item)
    _spacer(doc)

    # ── Section 7: Detailed Steps (dominant section) ──
    _h1(doc, "7.", "Detailed Procedure Steps")
    for phase in draft.phases:
        _h2(doc, f"{phase.phase_id}  {phase.title}")
        _body(doc, phase.objective, indent_cm=0.3)

        step_tbl = _make_table(doc, [
            "Step", "Actor", "Action", "System / Tool",
            "Output", "Evidence", "SLA"
        ])
        for i, step in enumerate(phase.steps):
            _add_row(step_tbl, [
                step.step_id, step.actor, step.action,
                step.system, step.output, step.evidence, step.timing,
            ], alt=(i % 2 == 0))
        _spacer(doc)

    # ── Section 8: Decisions / Exceptions / Escalation ──
    _h1(doc, "8.", "Decision Points, Exceptions and Escalations")
    if draft.decision_points:
        _h2(doc, "Decision Points")
        for dp in draft.decision_points:
            _bullet(doc, dp)
    _h2(doc, "Exceptions")
    _body(doc, draft.exceptions)
    _h2(doc, "Escalation Path")
    _body(doc, draft.escalation)
    _spacer(doc)

    # ── Section 9: Outputs / Records / Evidence ──
    _h1(doc, "9.", "Outputs, Records, Evidence and Forms")
    rec_tbl = _make_table(doc,
        ["Record / Evidence", "Owner", "Retention", "Storage Location"])
    for i, r in enumerate(draft.records):
        _add_row(rec_tbl, [r, "Cybersecurity Department",
                            "3 years", "Document Management System"],
                 alt=(i % 2 == 0))
    _spacer(doc)

    # ── Section 10: Time Controls ──
    _h1(doc, "10.", "Time Controls, SLAs and Control Checkpoints")
    tc_tbl = _make_table(doc, ["Control Checkpoint / SLA", "Requirement"])
    for i, tc in enumerate(draft.time_controls):
        _add_row(tc_tbl, [f"SLA {i+1}", tc], alt=(i % 2 == 0))
    _spacer(doc)

    # ── Section 11: Related Documents ──
    _h1(doc, "11.", "Related Documents")
    rd_tbl = _make_table(doc, ["Document", "Type"])
    for i, rd in enumerate(draft.related_docs):
        doc_type = ("Policy" if "POL-" in rd else
                    "Standard" if "STD-" in rd else
                    "Procedure" if "PRC-" in rd else "Reference")
        _add_row(rd_tbl, [rd, doc_type], alt=(i % 2 == 0))
    _spacer(doc)

    # ── Sections 12-15 ──
    _h1(doc, "12.", "Effective Date")
    _body(doc, draft.effective_date)
    _spacer(doc)

    _h1(doc, "13.", "Procedure Review")
    _body(doc, draft.review)
    _spacer(doc)

    _h1(doc, "14.", "Review and Approval")
    _approval_table(doc)

    _h1(doc, "15.", "Version Control")
    _version_table(doc, draft.version)

    path = os.path.join(output_dir, f"{draft.doc_id}_big4.docx")
    doc.save(path)
    return path
