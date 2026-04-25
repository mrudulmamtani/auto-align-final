"""
Ministry Renderer — converts ministry draft models to branded DOCX files.

Visual system:
  - RTL Arabic layout
  - Dark green headers (#0F4D43)
  - "Draft" watermark (large diagonal, light gray)
  - Classification footer on every internal page
  - Minimal English — only technical glosses in parentheses
"""
from __future__ import annotations
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from .ministry_models import (
    MinistryPolicyDraft, MinistryStandardDraft, MinistryProcedureDraft,
    DefinitionEntry, ApprovalStage, VersionRow, PolicyClause,
    StandardDomainCluster, ProcedurePhase,
)

# ── Colors ─────────────────────────────────────────────────────────────────────
PRIMARY_GREEN  = "0F4D43"
SECONDARY_GREEN= "2E7D6B"
NEUTRAL_LIGHT  = "F4F4F4"
BORDER_GRAY    = "D9D9D9"
WHITE          = "FFFFFF"


# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def _set_rtl_para(para):
    """Set paragraph to RTL bidirectional text."""
    pPr = para._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    bidi.set(qn("w:val"), "1")
    pPr.append(bidi)
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "right")
    pPr.append(jc)
    return para


def _set_rtl_run(run):
    """Set run to RTL."""
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rtl.set(qn("w:val"), "1")
    rPr.append(rtl)
    cs = OxmlElement("w:cs")
    rPr.append(cs)
    return run


def _set_cell_rtl(cell):
    for para in cell.paragraphs:
        _set_rtl_para(para)


def _add_watermark(doc: Document, text: str = "Draft"):
    """Add a diagonal watermark to the document header."""
    for section in doc.sections:
        header = section.header
        if not header.paragraphs:
            header.add_paragraph()
        para = header.paragraphs[0]
        run = para.add_run()
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(72)
        run.bold = True
        run.text = text
        pPr = para._p.get_or_add_pPr()
        jc = OxmlElement("w:jc")
        jc.set(qn("w:val"), "center")
        pPr.append(jc)


def _add_classification_footer(doc: Document, classification: str):
    """Add classification label to footer of all sections."""
    for section in doc.sections:
        footer = section.footer
        if footer.paragraphs:
            fp = footer.paragraphs[0]
        else:
            fp = footer.add_paragraph()
        fp.clear()
        _set_rtl_para(fp)
        run = fp.add_run(f"{classification}  |  Confidential — Internal Use Only")
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
        run.bold = True
        _set_rtl_run(run)

        # Page number on left side
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_page_number(doc: Document):
    """Add page number field to footers."""
    for section in doc.sections:
        footer = section.footer
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[-1] if len(footer.paragraphs) > 1 else footer.paragraphs[0]
        run = fp.add_run()
        fldChar = OxmlElement("w:fldChar")
        fldChar.set(qn("w:fldCharType"), "begin")
        instrText = OxmlElement("w:instrText")
        instrText.text = "PAGE"
        fldChar2 = OxmlElement("w:fldChar")
        fldChar2.set(qn("w:fldCharType"), "end")
        run._r.append(fldChar)
        run._r.append(instrText)
        run._r.append(fldChar2)


def _set_table_borders(table):
    """Apply thin gray grid borders to a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), BORDER_GRAY)
        tblBorders.append(border)
    tblPr.append(tblBorders)


def _green_header_row(table, texts: list[str]):
    """Style the first row of a table with dark green header."""
    row = table.rows[0]
    for i, cell in enumerate(row.cells):
        _set_cell_bg(cell, PRIMARY_GREEN)
        for para in cell.paragraphs:
            para.clear()
            _set_rtl_para(para)
            run = para.add_run(texts[i] if i < len(texts) else "")
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            _set_rtl_run(run)


def _add_section_heading(doc: Document, text: str, level: int = 1):
    """Add a section heading with ministry style."""
    para = doc.add_paragraph()
    _set_rtl_para(para)
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(14 if level == 1 else 12)
    run.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
    _set_rtl_run(run)
    return para


def _add_rtl_para(doc: Document, text: str, bold: bool = False, size: int = 11) -> None:
    """Add a RTL paragraph."""
    para = doc.add_paragraph()
    _set_rtl_para(para)
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    _set_rtl_run(run)


def _add_definitions_table(doc: Document, definitions: list[DefinitionEntry]):
    """Render a definition table with green header."""
    table = doc.add_table(rows=1 + len(definitions), cols=2)
    table.style = "Table Grid"
    _set_table_borders(table)
    _green_header_row(table, ["التعريف", "المصطلح"])
    for i, d in enumerate(definitions):
        row = table.rows[i + 1]
        def_cell = row.cells[0]
        term_cell = row.cells[1]
        # Term cell (right column in RTL display)
        _set_cell_rtl(term_cell)
        term_cell.paragraphs[0].clear()
        _set_rtl_para(term_cell.paragraphs[0])
        tr = term_cell.paragraphs[0].add_run(
            f"{d.term_ar}" + (f" ({d.term_en})" if d.term_en else "")
        )
        tr.bold = True
        tr.font.size = Pt(10)
        _set_rtl_run(tr)
        # Definition cell
        _set_cell_rtl(def_cell)
        def_cell.paragraphs[0].clear()
        _set_rtl_para(def_cell.paragraphs[0])
        dr = def_cell.paragraphs[0].add_run(d.definition_ar)
        dr.font.size = Pt(10)
        _set_rtl_run(dr)


def _add_approval_table(doc: Document, stages: list[ApprovalStage]):
    """Render the approval blocks table."""
    table = doc.add_table(rows=1 + len(stages), cols=4)
    table.style = "Table Grid"
    _set_table_borders(table)
    _green_header_row(table, ["التاريخ", "الاسم", "المنصب / الدور", "المرحلة"])
    for i, s in enumerate(stages):
        row = table.rows[i + 1]
        for j, text in enumerate([s.date_ar, s.name_ar, s.role_ar, s.stage_ar]):
            cell = row.cells[j]
            cell.paragraphs[0].clear()
            _set_rtl_para(cell.paragraphs[0])
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(10)
            _set_rtl_run(r)


def _add_version_table(doc: Document, rows: list[VersionRow]):
    """Render the version control table."""
    headers = ["تاريخ الاعتماد", "المُحدِّث", "ملخص التحديث", "نوع التحديث", "الإصدار"]
    table = doc.add_table(rows=1 + len(rows), cols=5)
    table.style = "Table Grid"
    _set_table_borders(table)
    _green_header_row(table, headers)
    for i, row in enumerate(rows):
        tr = table.rows[i + 1]
        for j, text in enumerate([row.approval_date, row.updated_by_ar, row.summary_ar, row.update_type_ar, row.version]):
            cell = tr.cells[j]
            cell.paragraphs[0].clear()
            _set_rtl_para(cell.paragraphs[0])
            r = cell.paragraphs[0].add_run(text)
            r.font.size = Pt(10)
            _set_rtl_run(r)


def _new_doc_with_branding(classification: str) -> Document:
    """Create a new Document with branding applied."""
    doc = Document()
    # Set default font to Arabic-compatible
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)
    # Set page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    _add_watermark(doc)
    _add_classification_footer(doc, classification)
    return doc


# ══════════════════════════════════════════════════════════════════════════════
# POLICY RENDERER
# ══════════════════════════════════════════════════════════════════════════════

def render_ministry_policy(draft: MinistryPolicyDraft, output_dir: str) -> str:
    """Render a MinistryPolicyDraft to a DOCX file. Returns the output path."""
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc_with_branding(draft.meta.classification)

    # ── Cover page ────────────────────────────────────────────────────────────
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_para(cover)
    cr = cover.add_run(draft.meta.title_ar)
    cr.bold = True
    cr.font.size = Pt(24)
    cr.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
    _set_rtl_run(cr)

    if draft.meta.title_en:
        ep = doc.add_paragraph()
        ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
        er = ep.add_run(draft.meta.title_en)
        er.italic = True
        er.font.size = Pt(16)

    doc.add_paragraph()
    meta_lines = [
        f"المعرّف: {draft.meta.doc_id}",
        f"الإصدار: {draft.meta.version}",
        f"التصنيف: {draft.meta.classification}",
        f"الجهة المسؤولة: {draft.meta.owner}",
    ]
    for line in meta_lines:
        mp = doc.add_paragraph()
        _set_rtl_para(mp)
        mr = mp.add_run(line)
        mr.font.size = Pt(11)
        _set_rtl_run(mr)

    doc.add_page_break()

    # ── Table of Contents placeholder ─────────────────────────────────────────
    _add_section_heading(doc, "جدول المحتويات")
    toc_items = [
        "1. التعريفات",
        "2. هدف السياسة",
        "3. نطاق السياسة",
        "4. بنود السياسة",
        "5. الأدوار والمسؤوليات",
        "6. الوثائق ذات الصلة",
        "7. تاريخ السريان",
        "8. مراجعة السياسة",
        "9. المراجعة والاعتماد",
        "10. سجل الإصدارات",
    ]
    for item in toc_items:
        ip = doc.add_paragraph()
        _set_rtl_para(ip)
        ir = ip.add_run(item)
        ir.font.size = Pt(11)
        _set_rtl_run(ir)

    doc.add_page_break()

    # ── Document Information ──────────────────────────────────────────────────
    _add_section_heading(doc, "معلومات الوثيقة")
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = "Table Grid"
    _set_table_borders(info_table)
    info_data = [
        ("عنوان الوثيقة", draft.meta.title_ar),
        ("معرّف الوثيقة", draft.meta.doc_id),
        ("الإصدار", draft.meta.version),
        ("الجهة المسؤولة", draft.meta.owner),
        ("التصنيف", draft.meta.classification),
        ("تاريخ الإصدار", draft.meta.issue_date or "[تاريخ الاعتماد]"),
    ]
    for i, (label, value) in enumerate(info_data):
        lc = info_table.rows[i].cells[1]
        vc = info_table.rows[i].cells[0]
        _set_cell_rtl(lc)
        lc.paragraphs[0].clear()
        _set_rtl_para(lc.paragraphs[0])
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        _set_rtl_run(lr)
        vc.paragraphs[0].clear()
        _set_rtl_para(vc.paragraphs[0])
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)
        _set_rtl_run(vr)

    doc.add_paragraph()

    # ── Section 1: Definitions ────────────────────────────────────────────────
    _add_section_heading(doc, "1. التعريفات")
    _add_definitions_table(doc, draft.definitions)
    doc.add_paragraph()

    # ── Section 2: Objective ──────────────────────────────────────────────────
    _add_section_heading(doc, "2. هدف السياسة")
    _add_rtl_para(doc, draft.objective_ar)
    doc.add_paragraph()

    # ── Section 3: Scope ──────────────────────────────────────────────────────
    _add_section_heading(doc, "3. نطاق السياسة")
    _add_rtl_para(doc, draft.scope_ar)
    doc.add_paragraph()

    # ── Section 4: Policy Statement (dominant) ────────────────────────────────
    _add_section_heading(doc, "4. بنود السياسة")
    _add_section_heading(doc, "أحكام عامة", level=2)
    for clause in draft.policy_clauses:
        cp = doc.add_paragraph()
        _set_rtl_para(cp)
        num_run = cp.add_run(f"{clause.clause_no}. ")
        num_run.bold = True
        num_run.font.size = Pt(11)
        num_run.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
        _set_rtl_run(num_run)
        txt_run = cp.add_run(clause.text_ar)
        txt_run.font.size = Pt(11)
        _set_rtl_run(txt_run)
        for bullet in clause.sub_bullets_ar:
            bp = doc.add_paragraph()
            _set_rtl_para(bp)
            pPr = bp._p.get_or_add_pPr()
            ind = OxmlElement("w:ind")
            ind.set(qn("w:right"), "720")
            pPr.append(ind)
            br = bp.add_run(f"• {bullet}")
            br.font.size = Pt(10)
            _set_rtl_run(br)
    doc.add_paragraph()

    # ── Section 5: Roles ──────────────────────────────────────────────────────
    _add_section_heading(doc, "5. الأدوار والمسؤوليات")
    _add_rtl_para(doc, draft.roles_ar)
    doc.add_paragraph()

    # ── Section 6: Related Documents ─────────────────────────────────────────
    _add_section_heading(doc, "6. الوثائق ذات الصلة")
    if draft.related_docs:
        rel_table = doc.add_table(rows=1 + len(draft.related_docs), cols=2)
        rel_table.style = "Table Grid"
        _set_table_borders(rel_table)
        _green_header_row(rel_table, ["النوع", "عنوان الوثيقة"])
        for i, rd in enumerate(draft.related_docs):
            row = rel_table.rows[i + 1]
            tc = row.cells[1]
            nc = row.cells[0]
            for cell, text in [(tc, rd.title_ar), (nc, rd.ref_type)]:
                cell.paragraphs[0].clear()
                _set_rtl_para(cell.paragraphs[0])
                r = cell.paragraphs[0].add_run(text)
                r.font.size = Pt(10)
                _set_rtl_run(r)
    doc.add_paragraph()

    # ── Section 7: Effective Date ─────────────────────────────────────────────
    _add_section_heading(doc, "7. تاريخ السريان")
    _add_rtl_para(doc, draft.effective_date_ar)
    doc.add_paragraph()

    # ── Section 8: Policy Review ──────────────────────────────────────────────
    _add_section_heading(doc, "8. مراجعة السياسة")
    _add_rtl_para(doc, draft.review_ar)
    doc.add_paragraph()

    # ── Section 9: Review and Approval ───────────────────────────────────────
    _add_section_heading(doc, "9. المراجعة والاعتماد")
    _add_approval_table(doc, draft.approval_stages)
    doc.add_paragraph()

    # ── Section 10: Version Control ───────────────────────────────────────────
    _add_section_heading(doc, "10. سجل الإصدارات")
    _add_version_table(doc, draft.version_rows)

    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}.docx")
    doc.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# STANDARD RENDERER
# ══════════════════════════════════════════════════════════════════════════════

def render_ministry_standard(draft: MinistryStandardDraft, output_dir: str) -> str:
    """Render a MinistryStandardDraft to DOCX. Returns the output path."""
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc_with_branding(draft.meta.classification)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_para(cover)
    cr = cover.add_run(draft.meta.title_ar)
    cr.bold = True
    cr.font.size = Pt(24)
    cr.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
    _set_rtl_run(cr)

    meta_lines = [
        f"المعرّف: {draft.meta.doc_id}",
        f"الإصدار: {draft.meta.version}",
        f"التصنيف: {draft.meta.classification}",
    ]
    for line in meta_lines:
        mp = doc.add_paragraph()
        _set_rtl_para(mp)
        _set_rtl_run(mp.add_run(line)).font.size = Pt(11)
    doc.add_page_break()

    # Notice page
    _add_section_heading(doc, "إشعار الملكية")
    _add_rtl_para(doc, draft.notice_ar)
    if draft.notice_en:
        p = doc.add_paragraph()
        r = p.add_run(draft.notice_en)
        r.italic = True
        r.font.size = Pt(10)
    doc.add_page_break()

    # Document information table
    _add_section_heading(doc, "معلومات الوثيقة")
    info_data = [
        ("عنوان الوثيقة", draft.meta.title_ar),
        ("معرّف الوثيقة", draft.meta.doc_id),
        ("الإصدار", draft.meta.version),
        ("الجهة المسؤولة", draft.meta.owner),
        ("التصنيف", draft.meta.classification),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.style = "Table Grid"
    _set_table_borders(info_table)
    for i, (label, value) in enumerate(info_data):
        lc = info_table.rows[i].cells[1]
        vc = info_table.rows[i].cells[0]
        lc.paragraphs[0].clear()
        _set_rtl_para(lc.paragraphs[0])
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        _set_rtl_run(lr)
        vc.paragraphs[0].clear()
        _set_rtl_para(vc.paragraphs[0])
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)
        _set_rtl_run(vr)
    doc.add_paragraph()

    # Approval table
    _add_section_heading(doc, "جدول الاعتماد")
    _add_approval_table(doc, draft.approval_stages)
    doc.add_page_break()

    # ToC
    _add_section_heading(doc, "جدول المحتويات")
    toc_items = [
        "1. الاختصارات والتعريفات",
        "2. الهدف",
        "3. نطاق التطبيق والانطباق",
        "4. المعايير",
        "5. الاستثناءات",
        "6. الأدوار والمسؤوليات",
        "7. التحديث والمراجعة",
        "8. الامتثال للمعيار",
    ]
    for item in toc_items:
        ip = doc.add_paragraph()
        _set_rtl_para(ip)
        _set_rtl_run(ip.add_run(item)).font.size = Pt(11)
    doc.add_page_break()

    # Section 1: Definitions
    _add_section_heading(doc, "1. الاختصارات والتعريفات")
    _add_definitions_table(doc, draft.definitions)
    doc.add_paragraph()

    # Section 2: Objective
    _add_section_heading(doc, "2. الهدف")
    _add_rtl_para(doc, draft.objective_ar)
    doc.add_paragraph()

    # Section 3: Scope
    _add_section_heading(doc, "3. نطاق التطبيق والانطباق")
    _add_rtl_para(doc, draft.scope_intro_ar)
    scope_cats = [
        ("الأنظمة والتطبيقات", draft.scope_systems_ar),
        ("الأدوار والصلاحيات", draft.scope_roles_ar),
        ("العمليات والمعاملات الحرجة", draft.scope_processes_ar),
        ("الأشخاص والجهات المشمولة", draft.scope_persons_ar),
    ]
    for cat_name, items in scope_cats:
        if items:
            _add_section_heading(doc, cat_name, level=2)
            for item in items:
                bp = doc.add_paragraph()
                _set_rtl_para(bp)
                _set_rtl_run(bp.add_run(f"• {item}")).font.size = Pt(10)
    doc.add_paragraph()

    # Section 4: Standards (dominant)
    _add_section_heading(doc, "4. المعايير")
    for cluster in draft.domain_clusters:
        _add_section_heading(doc, f"{cluster.cluster_id} {cluster.title_ar}", level=2)
        _add_rtl_para(doc, f"الهدف: {cluster.objective_ar}")
        _add_rtl_para(doc, f"المخاطر المحتملة: {cluster.potential_risks_ar}")
        for clause in cluster.clauses:
            cp = doc.add_paragraph()
            _set_rtl_para(cp)
            num_r = cp.add_run(f"{clause.clause_id}  ")
            num_r.bold = True
            num_r.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
            _set_rtl_run(num_r)
            txt_r = cp.add_run(clause.text_ar)
            txt_r.font.size = Pt(11)
            _set_rtl_run(txt_r)
            if clause.guidance_ar:
                gp = doc.add_paragraph()
                _set_rtl_para(gp)
                pPr = gp._p.get_or_add_pPr()
                ind = OxmlElement("w:ind")
                ind.set(qn("w:right"), "720")
                pPr.append(ind)
                _set_rtl_run(gp.add_run(f"توجيه: {clause.guidance_ar}")).font.size = Pt(10)
        doc.add_paragraph()

    # Section 5: Exceptions
    _add_section_heading(doc, "5. الاستثناءات")
    _add_rtl_para(doc, draft.exceptions_ar)
    doc.add_paragraph()

    # Section 6: Roles
    _add_section_heading(doc, "6. الأدوار والمسؤوليات")
    if draft.roles_responsibilities:
        roles_table = doc.add_table(rows=1 + len(draft.roles_responsibilities), cols=2)
        roles_table.style = "Table Grid"
        _set_table_borders(roles_table)
        _green_header_row(roles_table, ["المسؤوليات", "الدور"])
        for i, (role, duties) in enumerate(draft.roles_responsibilities.items()):
            row = roles_table.rows[i + 1]
            for j, text in enumerate([duties, role]):
                c = row.cells[j]
                c.paragraphs[0].clear()
                _set_rtl_para(c.paragraphs[0])
                r = c.paragraphs[0].add_run(text)
                r.font.size = Pt(10)
                _set_rtl_run(r)
    doc.add_paragraph()

    # Section 7: Update and Review
    _add_section_heading(doc, "7. التحديث والمراجعة")
    _add_rtl_para(doc, draft.update_review_ar)
    doc.add_paragraph()

    # Section 8: Compliance
    _add_section_heading(doc, "8. الامتثال للمعيار")
    _add_rtl_para(doc, draft.compliance_ar)
    doc.add_paragraph()

    # Version control
    _add_section_heading(doc, "سجل الإصدارات")
    _add_version_table(doc, draft.version_rows)

    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}.docx")
    doc.save(out_path)
    return out_path


# ══════════════════════════════════════════════════════════════════════════════
# PROCEDURE RENDERER
# ══════════════════════════════════════════════════════════════════════════════

def render_ministry_procedure(draft: MinistryProcedureDraft, output_dir: str) -> str:
    """Render a MinistryProcedureDraft to DOCX. Returns the output path."""
    os.makedirs(output_dir, exist_ok=True)
    doc = _new_doc_with_branding(draft.meta.classification)

    # Cover
    cover = doc.add_paragraph()
    cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _set_rtl_para(cover)
    cr = cover.add_run(draft.meta.title_ar)
    cr.bold = True
    cr.font.size = Pt(24)
    cr.font.color.rgb = RGBColor(0x0F, 0x4D, 0x43)
    _set_rtl_run(cr)
    meta_lines = [
        f"المعرّف: {draft.meta.doc_id}",
        f"الإصدار: {draft.meta.version}",
        f"التصنيف: {draft.meta.classification}",
    ]
    for line in meta_lines:
        mp = doc.add_paragraph()
        _set_rtl_para(mp)
        _set_rtl_run(mp.add_run(line)).font.size = Pt(11)
    doc.add_page_break()

    # Document info
    _add_section_heading(doc, "معلومات الوثيقة")
    info_data = [
        ("عنوان الإجراء", draft.meta.title_ar),
        ("معرّف الوثيقة", draft.meta.doc_id),
        ("السياسة الأم", draft.parent_policy_id or "-"),
        ("المعيار الأم", draft.parent_standard_id or "-"),
        ("الإصدار", draft.meta.version),
        ("الجهة المسؤولة", draft.meta.owner),
        ("التصنيف", draft.meta.classification),
    ]
    info_table = doc.add_table(rows=len(info_data), cols=2)
    info_table.style = "Table Grid"
    _set_table_borders(info_table)
    for i, (label, value) in enumerate(info_data):
        lc = info_table.rows[i].cells[1]
        vc = info_table.rows[i].cells[0]
        lc.paragraphs[0].clear()
        _set_rtl_para(lc.paragraphs[0])
        lr = lc.paragraphs[0].add_run(label)
        lr.bold = True
        lr.font.size = Pt(10)
        _set_rtl_run(lr)
        vc.paragraphs[0].clear()
        _set_rtl_para(vc.paragraphs[0])
        vr = vc.paragraphs[0].add_run(value)
        vr.font.size = Pt(10)
        _set_rtl_run(vr)
    doc.add_paragraph()

    # ToC
    _add_section_heading(doc, "جدول المحتويات")
    toc_items = [
        "1. التعريفات والاختصارات",
        "2. هدف الإجراء",
        "3. النطاق والانطباق",
        "4. الأدوار والمسؤوليات",
        "5. نظرة عامة على الإجراء",
        "6. المحفّزات والمتطلبات المسبقة والمدخلات والأدوات",
        "7. خطوات الإجراء التفصيلية",
        "8. نقاط القرار والاستثناءات والتصعيد",
        "9. المخرجات والسجلات والأدلة والنماذج",
        "10. ضوابط الوقت ومستويات الخدمة ونقاط التحقق",
        "11. الوثائق ذات الصلة",
        "12. تاريخ السريان",
        "13. مراجعة الإجراء",
        "14. المراجعة والاعتماد",
        "15. سجل الإصدارات",
    ]
    for item in toc_items:
        ip = doc.add_paragraph()
        _set_rtl_para(ip)
        _set_rtl_run(ip.add_run(item)).font.size = Pt(11)
    doc.add_page_break()

    # Section 1
    _add_section_heading(doc, "1. التعريفات والاختصارات")
    _add_definitions_table(doc, draft.definitions)
    doc.add_paragraph()

    # Section 2
    _add_section_heading(doc, "2. هدف الإجراء")
    _add_rtl_para(doc, draft.objective_ar)
    doc.add_paragraph()

    # Section 3
    _add_section_heading(doc, "3. النطاق والانطباق")
    _add_rtl_para(doc, draft.scope_ar)
    doc.add_paragraph()

    # Section 4: Roles
    _add_section_heading(doc, "4. الأدوار والمسؤوليات")
    for role_item in draft.roles:
        _add_rtl_para(doc, role_item.role_ar, bold=True)
        for resp in role_item.responsibilities_ar:
            bp = doc.add_paragraph()
            _set_rtl_para(bp)
            _set_rtl_run(bp.add_run(f"• {resp}")).font.size = Pt(10)
    doc.add_paragraph()

    # Section 5: Overview
    _add_section_heading(doc, "5. نظرة عامة على الإجراء")
    _add_rtl_para(doc, draft.overview_ar)
    doc.add_paragraph()

    # Section 6: Triggers, prerequisites, inputs, tools
    _add_section_heading(doc, "6. المحفّزات والمتطلبات المسبقة والمدخلات والأدوات")
    for label, items in [
        ("المحفّزات", draft.triggers_ar),
        ("المتطلبات المسبقة", draft.prerequisites_ar),
        ("المدخلات", draft.inputs_ar),
        ("الأدوات والأنظمة", draft.tools_ar),
    ]:
        if items:
            _add_rtl_para(doc, label, bold=True)
            for item in items:
                bp = doc.add_paragraph()
                _set_rtl_para(bp)
                _set_rtl_run(bp.add_run(f"• {item}")).font.size = Pt(10)
    doc.add_paragraph()

    # Section 7: Detailed procedure steps (dominant)
    _add_section_heading(doc, "7. خطوات الإجراء التفصيلية")
    for phase in draft.phases:
        _add_section_heading(doc, f"{phase.phase_id} {phase.phase_title_ar}", level=2)
        if phase.phase_objective_ar:
            _add_rtl_para(doc, phase.phase_objective_ar)
        # Steps table
        if phase.steps:
            steps_table = doc.add_table(rows=1 + len(phase.steps), cols=6)
            steps_table.style = "Table Grid"
            _set_table_borders(steps_table)
            _green_header_row(steps_table, ["الدليل", "المخرج", "النظام", "الإجراء", "المنفّذ", "رقم الخطوة"])
            for i, step in enumerate(phase.steps):
                row = steps_table.rows[i + 1]
                for j, text in enumerate([step.evidence_ar, step.output_ar, step.system_ar, step.action_ar, step.actor_ar, step.step_id]):
                    c = row.cells[j]
                    c.paragraphs[0].clear()
                    _set_rtl_para(c.paragraphs[0])
                    r = c.paragraphs[0].add_run(text)
                    r.font.size = Pt(9)
                    _set_rtl_run(r)
        doc.add_paragraph()

    # Section 8
    _add_section_heading(doc, "8. نقاط القرار والاستثناءات والتصعيد")
    if draft.decision_points_ar:
        _add_rtl_para(doc, "نقاط القرار:", bold=True)
        for dp in draft.decision_points_ar:
            bp = doc.add_paragraph()
            _set_rtl_para(bp)
            _set_rtl_run(bp.add_run(f"• {dp}")).font.size = Pt(10)
    _add_rtl_para(doc, "الاستثناءات:", bold=True)
    _add_rtl_para(doc, draft.exceptions_ar)
    _add_rtl_para(doc, "التصعيد:", bold=True)
    _add_rtl_para(doc, draft.escalation_ar)
    doc.add_paragraph()

    # Section 9
    _add_section_heading(doc, "9. المخرجات والسجلات والأدلة والنماذج")
    for label, items in [
        ("المخرجات", draft.outputs_ar),
        ("السجلات", draft.records_ar),
        ("الأدلة", draft.evidence_ar),
        ("النماذج", draft.forms_ar),
    ]:
        if items:
            _add_rtl_para(doc, label, bold=True)
            for item in items:
                bp = doc.add_paragraph()
                _set_rtl_para(bp)
                _set_rtl_run(bp.add_run(f"• {item}")).font.size = Pt(10)
    doc.add_paragraph()

    # Section 10
    _add_section_heading(doc, "10. ضوابط الوقت ومستويات الخدمة ونقاط التحقق")
    for tc in draft.time_controls_ar:
        bp = doc.add_paragraph()
        _set_rtl_para(bp)
        _set_rtl_run(bp.add_run(f"• {tc}")).font.size = Pt(10)
    doc.add_paragraph()

    # Section 11: Related docs
    _add_section_heading(doc, "11. الوثائق ذات الصلة")
    if draft.related_docs:
        rel_table = doc.add_table(rows=1 + len(draft.related_docs), cols=2)
        rel_table.style = "Table Grid"
        _set_table_borders(rel_table)
        _green_header_row(rel_table, ["النوع", "عنوان الوثيقة"])
        for i, rd in enumerate(draft.related_docs):
            row = rel_table.rows[i + 1]
            for j, text in enumerate([rd.ref_type, rd.title_ar]):
                c = row.cells[j]
                c.paragraphs[0].clear()
                _set_rtl_para(c.paragraphs[0])
                _set_rtl_run(c.paragraphs[0].add_run(text)).font.size = Pt(10)
    doc.add_paragraph()

    # Section 12
    _add_section_heading(doc, "12. تاريخ السريان")
    _add_rtl_para(doc, draft.effective_date_ar)
    doc.add_paragraph()

    # Section 13
    _add_section_heading(doc, "13. مراجعة الإجراء")
    _add_rtl_para(doc, draft.review_ar)
    doc.add_paragraph()

    # Section 14
    _add_section_heading(doc, "14. المراجعة والاعتماد")
    _add_approval_table(doc, draft.approval_stages)
    doc.add_paragraph()

    # Section 15
    _add_section_heading(doc, "15. سجل الإصدارات")
    _add_version_table(doc, draft.version_rows)

    out_path = os.path.join(output_dir, f"{draft.meta.doc_id}.docx")
    doc.save(out_path)
    return out_path
